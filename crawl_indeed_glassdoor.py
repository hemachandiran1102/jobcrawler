"""
crawl_indeed_glassdoor.py — Indeed & Glassdoor Job Crawler -> Excel, CSV & Word
================================================================================
3-Tier API-First Architecture (mirroring LinkedIn Voyager approach):

  Tier 1 — API Layer:
    • Indeed: RSS feeds (existing) + hidden JSON endpoints
    • Glassdoor: Internal GraphQL API (JobSearchResultsQuery)

  Tier 2 — Google Search Fallback:
    • site:indeed.com / site:glassdoor.com search parsing
    • No browser needed, uses requests with rotating User-Agents

  Tier 3 — Browser with XHR Interception:
    • Playwright captures AJAX/JSON responses instead of fragile DOM selectors
    • Anti-detection stealth patterns
    • DOM scraping as absolute last resort

Features:
- Multi-tier scraping with automatic fallback cascade
- Complete data schema: Job Title, Company, Location, Country, Match Score,
  Visa Sponsorship, Required Skills, Tailored Resume, Source, Job URL,
  Posted Date, Crawl Date, Applied Status.
- Multi-sheet Excel workbook (All Jobs master sheet + Date Tabs)
- Single Master CSV with automatic deduplication
- Word document report generation

Usage:
  python crawl_indeed_glassdoor.py --days 2
  python crawl_indeed_glassdoor.py --countries london toronto dubai amsterdam --days 2
  python crawl_indeed_glassdoor.py --sources indeed glassdoor --days 2
  python crawl_indeed_glassdoor.py --no-headless
  python crawl_indeed_glassdoor.py --skip-google          # skip Tier 2
  python crawl_indeed_glassdoor.py --skip-browser          # skip Tier 3
"""

import os
import sys
import re
import time
import json
import random
import hashlib
import argparse
import urllib.parse
import xml.etree.ElementTree as ET
from pathlib import Path
from datetime import datetime, timedelta
import requests
import pandas as pd

try:
    from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout
    HAS_PLAYWRIGHT = True
except ImportError:
    HAS_PLAYWRIGHT = False

try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False

# ══════════════════════════════════════════════════════════════════════
# PATHS & CONFIGURATION
# ══════════════════════════════════════════════════════════════════════
_SCRIPT_DIR = Path(__file__).parent.resolve()
WORK_DIR    = Path(os.environ.get("WORK_DIR", str(_SCRIPT_DIR)))
SESSION_DIR = WORK_DIR / ".cf_session"

timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
RUN_DIR   = WORK_DIR / f"crawl_indeed_glassdoor_{timestamp}"
CSV_PATH  = RUN_DIR / f"indeed_glassdoor_jobs_{timestamp}.csv"
DOCX_PATH = RUN_DIR / f"Indeed_Glassdoor_Jobs_{timestamp}.docx"
EXCEL_PATH = RUN_DIR / f"indeed_glassdoor_jobs_{timestamp}.xlsx"

# Single Master Spreadsheets
MASTER_CSV_PATH   = WORK_DIR / "indeed_glassdoor_jobs.csv"
MASTER_EXCEL_PATH = WORK_DIR / "indeed_glassdoor_jobs.xlsx"
CONFIG_PATH       = WORK_DIR / "google_sheets_config.json"

# Glassdoor CSRF cache
GD_CSRF_CACHE     = SESSION_DIR / "gd_csrf.json"


# ══════════════════════════════════════════════════════════════════════
# TARGET COUNTRIES (All 22 Markets & Resumes)
# ══════════════════════════════════════════════════════════════════════
TARGET_COUNTRIES = [
    # ── Tier 1: Primary English-Speaking / High Feasibility / Major Tech Hubs ──
    {"name": "United Kingdom",       "flag": "🇬🇧", "tier": 1, "indeed_domain": "uk.indeed.com",
     "gd_location_id": 2, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_UNITED_KINGDOM.docx"},
    {"name": "Canada",               "flag": "🇨🇦", "tier": 1, "indeed_domain": "ca.indeed.com",
     "gd_location_id": 3, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_CANADA.docx"},
    {"name": "United Arab Emirates", "flag": "🇦🇪", "tier": 1, "indeed_domain": "ae.indeed.com",
     "gd_location_id": 247, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_UAE.docx"},
    {"name": "Saudi Arabia",         "flag": "🇸🇦", "tier": 1, "indeed_domain": "sa.indeed.com",
     "gd_location_id": 211, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_SAUDI_ARABIA.docx"},
    {"name": "Qatar",                "flag": "🇶🇦", "tier": 1, "indeed_domain": "qa.indeed.com",
     "gd_location_id": 202, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_QATAR.docx"},
    {"name": "Netherlands",          "flag": "🇳🇱", "tier": 1, "indeed_domain": "nl.indeed.com",
     "gd_location_id": 178, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_NETHERLANDS.docx"},
    {"name": "Ireland",              "flag": "🇮🇪", "tier": 1, "indeed_domain": "ie.indeed.com",
     "gd_location_id": 119, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_IRELAND.docx"},
    {"name": "Sweden",               "flag": "🇸🇪", "tier": 1, "indeed_domain": "se.indeed.com",
     "gd_location_id": 223, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_SWEDEN.docx"},
    {"name": "Denmark",              "flag": "🇩🇰", "tier": 1, "indeed_domain": "dk.indeed.com",
     "gd_location_id": 63, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_DENMARK.docx"},
    {"name": "Finland",              "flag": "🇫🇮", "tier": 1, "indeed_domain": "fi.indeed.com",
     "gd_location_id": 76, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_FINLAND.docx"},
    {"name": "Australia",            "flag": "🇦🇺", "tier": 1, "indeed_domain": "au.indeed.com",
     "gd_location_id": 13, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_AUSTRALIA.docx"},
    {"name": "Singapore",            "flag": "🇸🇬", "tier": 1, "indeed_domain": "sg.indeed.com",
     "gd_location_id": 215, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_SINGAPORE.docx"},
    {"name": "New Zealand",          "flag": "🇳🇿", "tier": 1, "indeed_domain": "nz.indeed.com",
     "gd_location_id": 182, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_NEW_ZEALAND.docx"},

    # ── Tier 2: Strong European & Arabian Tech Markets ──
    {"name": "Kuwait",               "flag": "🇰🇼", "tier": 2, "indeed_domain": "kw.indeed.com",
     "gd_location_id": 143, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_KUWAIT.docx"},
    {"name": "Bahrain",              "flag": "🇧🇭", "tier": 2, "indeed_domain": "bh.indeed.com",
     "gd_location_id": 19, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_BAHRAIN.docx"},
    {"name": "Oman",                 "flag": "🇴🇲", "tier": 2, "indeed_domain": "om.indeed.com",
     "gd_location_id": 189, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_OMAN.docx"},
    {"name": "France",               "flag": "🇫🇷", "tier": 2, "indeed_domain": "fr.indeed.com",
     "gd_location_id": 77, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_FRANCE.docx"},
    {"name": "Portugal",             "flag": "🇵🇹", "tier": 2, "indeed_domain": "pt.indeed.com",
     "gd_location_id": 200, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_PORTUGAL.docx"},
    {"name": "Poland",               "flag": "🇵🇱", "tier": 2, "indeed_domain": "pl.indeed.com",
     "gd_location_id": 198, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_POLAND.docx"},
    {"name": "Belgium",              "flag": "🇧🇪", "tier": 2, "indeed_domain": "be.indeed.com",
     "gd_location_id": 22, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_BELGIUM.docx"},
    {"name": "Austria",              "flag": "🇦🇹", "tier": 2, "indeed_domain": "at.indeed.com",
     "gd_location_id": 12, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_AUSTRIA.docx"},
    {"name": "Malaysia",             "flag": "🇲🇾", "tier": 2, "indeed_domain": "malaysia.indeed.com",
     "gd_location_id": 155, "gd_location_type": "N",
     "resume": "Country_Resumes/Hemachandiran_Giri_CV_MALAYSIA.docx"},
]

SEARCH_TERMS = [
    "DevOps Engineer",
    "Cloud Engineer",
    "Site Reliability Engineer",
    "Platform Engineer",
]

VISA_KEYWORDS = [
    "visa", "sponsorship", "relocation", "work permit", "blue card",
    "kennismigrant", "critical skills", "passeport talent", "tech visa",
    "red-white-red", "tier 2", "skilled worker", "immigration",
    "global skills", "express entry", "green visa", "golden visa", "iqama",
]

SKILL_TAGS = [
    "AWS", "Kubernetes", "Docker", "Terraform", "CI/CD", "Python", "Linux",
    "GCP", "Azure", "Ansible", "Golang", "Kafka", "Jenkins", "ArgoCD",
    "Helm", "GitOps", "Prometheus", "Grafana", "EKS", "ECS", "Vault",
    "Pulumi", "Bash", "GitHub Actions", "GitLab CI", "Datadog",
]

# Rotating User-Agent pool for anti-detection
USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 14_5) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:126.0) Gecko/20100101 Firefox/126.0",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 14.5; rv:126.0) Gecko/20100101 Firefox/126.0",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
    "Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:126.0) Gecko/20100101 Firefox/126.0",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36 Edg/126.0.0.0",
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 14_5) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.4.1 Safari/605.1.15",
]


# ══════════════════════════════════════════════════════════════════════
# UTILITIES & NORMALIZATION
# ══════════════════════════════════════════════════════════════════════
def log(msg, level="INFO"):
    ts   = datetime.now().strftime("%H:%M:%S")
    icon = {"INFO": "[i]", "OK": "[OK]", "WARN": "[!]", "ERROR": "[ERR]"}.get(level, "")
    safe = str(msg).encode("ascii", errors="replace").decode("ascii")
    print(f"[{ts}] {icon}  {safe}", flush=True)


def delay(lo=1.0, hi=2.0):
    time.sleep(random.uniform(lo, hi))


def random_ua() -> str:
    return random.choice(USER_AGENTS)


def normalize_url(url: str) -> str:
    """Extract canonical job URL for Indeed, Glassdoor, or other job boards."""
    u = str(url or "").strip()
    if not u or u.lower() in {"n/a", "none", "nan", "", "-"}:
        return ""

    # Indeed: https://www.indeed.com/viewjob?jk=1234567890abcdef
    m_indeed = re.search(r"[?&]jk=([a-zA-Z0-9]+)", u)
    if m_indeed:
        return f"https://www.indeed.com/viewjob?jk={m_indeed.group(1)}"

    # Glassdoor: https://www.glassdoor.com/job-listing/?jl=1234567890
    m_gd = re.search(r"(?:jl=|jobListingId=|job-listing/.*?jl=)(\d+)", u)
    if m_gd:
        return f"https://www.glassdoor.com/job-listing/?jl={m_gd.group(1)}"

    # LinkedIn: https://www.linkedin.com/jobs/view/1234567890
    m_li = re.search(r"/jobs/view/(?:[^\s/?#]*-)?(\d{6,14})", u)
    if m_li:
        return f"https://www.linkedin.com/jobs/view/{m_li.group(1)}"

    clean = u.split("#")[0]
    base = clean.split("?")[0].rstrip("/")
    if base.startswith("http://"):
        base = "https://" + base[7:]
    return base


def extract_skills(text: str) -> str:
    t = (text or "").lower()
    found = [s for s in SKILL_TAGS if s.lower() in t]
    return ", ".join(found[:10]) if found else "AWS, Kubernetes, CI/CD, Terraform, Linux"


def extract_visa(text: str) -> str:
    t = (text or "").lower()
    return "Yes" if any(k in t for k in VISA_KEYWORDS) else "No"


def calculate_match_score(title: str, desc: str) -> str:
    t = f"{title} {desc}".lower()
    score = 80
    if any(k in t for k in ["devops", "cloud", "reliability", "platform", "infrastructure", "kubernetes", "aws"]):
        score += 10
    if any(k in t for k in ["terraform", "ci/cd", "docker", "python", "linux", "gitops"]):
        score += 5
    if any(k in t for k in VISA_KEYWORDS):
        score += 5
    return f"{min(100, score)}%"


def normalise_posted_date(value) -> str:
    """Return a dashboard-friendly ISO date from various date formats."""
    raw = str(value or "").strip()
    if not raw or raw.lower() in {"n/a", "unknown", "past week", "past month"}:
        return "Unknown"

    iso = re.search(r"\b(20\d{2}-\d{2}-\d{2})", raw)
    if iso:
        return iso.group(1)

    text = raw.lower().replace("reposted", "").strip()
    now = datetime.now()
    if text in {"today", "just now", "new"}:
        return now.strftime("%Y-%m-%d")
    if text == "yesterday":
        return (now - timedelta(days=1)).strftime("%Y-%m-%d")

    relative = re.search(r"(\d+)\s*(minute|hour|day|week|month)s?\s+ago", text)
    if relative:
        amount, unit = int(relative.group(1)), relative.group(2)
        days_delta = {"minute": 0, "hour": 0, "day": amount, "week": amount * 7,
                      "month": amount * 30}[unit]
        return (now - timedelta(days=days_delta)).strftime("%Y-%m-%d")

    # Glassdoor format: "3d", "1w", "30d+"
    gd_relative = re.search(r"(\d+)\s*([dwm])\+?", text)
    if gd_relative:
        amount = int(gd_relative.group(1))
        unit = gd_relative.group(2)
        days_delta = {"d": amount, "w": amount * 7, "m": amount * 30}[unit]
        return (now - timedelta(days=days_delta)).strftime("%Y-%m-%d")

    return raw


def extract_city(location_str: str) -> str:
    loc = str(location_str or "").strip()
    if not loc or loc.lower() in {"n/a", "unknown", "remote", "global"}:
        return "Unknown"
    parts = [p.strip() for p in loc.split(",") if p.strip()]
    city_candidate = parts[0]
    city_candidate = re.sub(r"\b(Greater|Area|Metropolitan|Region|City of)\b", "", city_candidate, flags=re.IGNORECASE).strip()
    return city_candidate or "Unknown"


def filter_target_countries(targets: list, countries_filter) -> list:
    if not countries_filter:
        return targets

    alias_map = {
        "uk": ["United Kingdom"], "london": ["United Kingdom"], "manchester": ["United Kingdom"], "edinburgh": ["United Kingdom"],
        "canada": ["Canada"], "toronto": ["Canada"], "vancouver": ["Canada"], "montreal": ["Canada"], "ottawa": ["Canada"],
        "uae": ["United Arab Emirates"], "dubai": ["United Arab Emirates"], "abu dhabi": ["United Arab Emirates"],
        "saudi": ["Saudi Arabia"], "saudi arabia": ["Saudi Arabia"], "riyadh": ["Saudi Arabia"], "jeddah": ["Saudi Arabia"],
        "qatar": ["Qatar"], "doha": ["Qatar"],
        "netherlands": ["Netherlands"], "amsterdam": ["Netherlands"], "rotterdam": ["Netherlands"], "utrecht": ["Netherlands"],
        "ireland": ["Ireland"], "dublin": ["Ireland"], "cork": ["Ireland"],
        "france": ["France"], "paris": ["France"], "lyon": ["France"],
        "australia": ["Australia"], "sydney": ["Australia"], "melbourne": ["Australia"], "brisbane": ["Australia"],
        "singapore": ["Singapore"], "malaysia": ["Malaysia"], "kuala lumpur": ["Malaysia"],
        "new zealand": ["New Zealand"], "auckland": ["New Zealand"], "wellington": ["New Zealand"],
        "poland": ["Poland"], "warsaw": ["Poland"], "krakow": ["Poland"],
        "portugal": ["Portugal"], "lisbon": ["Portugal"], "porto": ["Portugal"],
        "germany": ["Germany"], "berlin": ["Germany"], "munich": ["Germany"],
        "sweden": ["Sweden"], "stockholm": ["Sweden"], "denmark": ["Denmark"], "copenhagen": ["Denmark"],
        "finland": ["Finland"], "helsinki": ["Finland"], "belgium": ["Belgium"], "brussels": ["Belgium"],
        "austria": ["Austria"], "vienna": ["Austria"], "kuwait": ["Kuwait"], "bahrain": ["Bahrain"], "oman": ["Oman"],
        "arabian": ["United Arab Emirates", "Saudi Arabia", "Qatar", "Kuwait", "Bahrain", "Oman"],
        "gcc": ["United Arab Emirates", "Saudi Arabia", "Qatar", "Kuwait", "Bahrain", "Oman"],
        "middle east": ["United Arab Emirates", "Saudi Arabia", "Qatar", "Kuwait", "Bahrain", "Oman"],
    }

    target_names = set()
    for f in countries_filter:
        fl = f.strip().lower()
        if fl in alias_map:
            for mapped in alias_map[fl]:
                target_names.add(mapped.lower())
        else:
            for c in targets:
                c_lower = c["name"].lower()
                if fl in c_lower or c_lower in fl:
                    target_names.add(c_lower)

    filtered = [c for c in targets if c["name"].lower() in target_names]
    return filtered or targets


def deduplicate_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    """Deduplicate dataframe based on normalized job URL, preserving metadata."""
    if df is None or df.empty:
        return pd.DataFrame()

    records = df.to_dict(orient="records")
    unique_map = {}
    dupes_removed = 0

    for rec in records:
        raw_url = str(rec.get("Job URL", "")).strip()
        norm_url = normalize_url(raw_url)

        if norm_url:
            key = f"URL:{norm_url}"
            rec["Job URL"] = norm_url
        else:
            comp = str(rec.get("Company", "")).strip().lower()
            title = str(rec.get("Job Title", "")).strip().lower()
            loc = str(rec.get("Location", rec.get("Country", ""))).strip().lower()
            key = f"TITLE:{comp}|{title}|{loc}"

        if key in unique_map:
            dupes_removed += 1
            existing = unique_map[key]
            # Preserve Applied Status
            if str(rec.get("Applied Status", "")).strip().lower() == "yes" or \
               str(rec.get("Applied", "")).strip().lower() == "yes":
                existing["Applied Status"] = "Yes"
                existing["Applied"] = "Yes"
            if rec.get("Notes") and not existing.get("Notes"):
                existing["Notes"] = rec["Notes"]
            existing_date = str(existing.get("Crawl Date", existing.get("Date", "")))
            rec_date = str(rec.get("Crawl Date", rec.get("Date", "")))
            if rec_date and (not existing_date or rec_date < existing_date):
                existing["Crawl Date"] = rec_date
                existing["Date"] = rec_date
            for k, v in rec.items():
                if v and (not existing.get(k) or str(existing.get(k)).strip() in {"", "N/A", "Unknown"}):
                    existing[k] = v
        else:
            unique_map[key] = rec

    if dupes_removed > 0:
        log(f"Removed {dupes_removed} duplicate job records.", "INFO")

    return pd.DataFrame(list(unique_map.values()))


def _make_job_record(title, company, location, country, keyword, posted_date,
                     source, job_url, desc="", easy_apply="No", remote=None):
    """Create a standardised job record dict (used by all tiers)."""
    full_text = f"{title} {desc} {location}"
    if remote is None:
        remote = "Remote" if "remote" in full_text.lower() else "On-site / Hybrid"
    return {
        "Job Title": title,
        "Company": company or "Company not stated",
        "Location": location or country["name"],
        "Country": country["name"],
        "Flag": country["flag"],
        "Tier": country["tier"],
        "Search Keyword": keyword,
        "Posted Date": normalise_posted_date(posted_date),
        "Crawl Date": datetime.now().strftime("%Y-%m-%d"),
        "Source": source,
        "Easy Apply": easy_apply,
        "Remote / Workplace": remote,
        "Match Score": calculate_match_score(title, desc),
        "Visa Sponsorship Mentioned": extract_visa(f"{title} {desc}"),
        "Required Skills": extract_skills(f"{title} {desc}"),
        "Resume File Path": country["resume"],
        "Applied Status": "No",
        "Notes": "",
        "Job URL": normalize_url(job_url),
    }


# ══════════════════════════════════════════════════════════════════════
# TIER 1A — INDEED API LAYER (RSS + Hidden JSON)
# ══════════════════════════════════════════════════════════════════════

def fetch_indeed_rss(country: dict, keyword: str, days: int = 2) -> list:
    """Fetch Indeed jobs via fast XML RSS endpoints."""
    domain = country.get("indeed_domain", "www.indeed.com")
    kw_enc = urllib.parse.quote_plus(keyword)
    loc_enc = urllib.parse.quote_plus(country["name"])

    url = f"https://{domain}/rss?q={kw_enc}&l={loc_enc}&fromage={days}&sort=date"
    headers = {
        "User-Agent": random_ua(),
        "Accept": "application/rss+xml, text/xml, application/xml;q=0.9, */*;q=0.8"
    }

    jobs = []
    try:
        resp = requests.get(url, headers=headers, timeout=15)
        if resp.status_code == 200 and resp.text:
            root = ET.fromstring(resp.content)
            items = root.findall(".//item")
            for item in items:
                title_elem = item.find("title")
                link_elem = item.find("link")
                desc_elem = item.find("description")
                pub_elem = item.find("pubDate")
                source_elem = item.find("source")

                title_raw = title_elem.text if title_elem is not None and title_elem.text else ""
                link_raw = link_elem.text if link_elem is not None and link_elem.text else ""
                desc_raw = desc_elem.text if desc_elem is not None and desc_elem.text else ""
                pub_raw = pub_elem.text if pub_elem is not None and pub_elem.text else ""
                comp_raw = source_elem.text if source_elem is not None and source_elem.text else ""

                if not comp_raw and " - " in title_raw:
                    parts = title_raw.rsplit(" - ", 1)
                    title_clean = parts[0].strip()
                    comp_raw = parts[1].strip()
                else:
                    title_clean = title_raw.strip()

                if not title_clean or not link_raw:
                    continue

                posted_date = datetime.now().strftime("%Y-%m-%d")
                if pub_raw:
                    try:
                        dt = pd.to_datetime(pub_raw)
                        posted_date = dt.strftime("%Y-%m-%d")
                    except Exception:
                        pass

                jobs.append(_make_job_record(
                    title=title_clean, company=comp_raw, location=country["name"],
                    country=country, keyword=keyword, posted_date=posted_date,
                    source="Indeed-RSS", job_url=link_raw, desc=desc_raw,
                    easy_apply="Yes" if "easy apply" in desc_raw.lower() else "No",
                ))
    except Exception as e:
        log(f"    Indeed RSS warning ({country['name']} - '{keyword}'): {e}", "WARN")

    return jobs


def fetch_indeed_json_api(country: dict, keyword: str, days: int = 2, max_pages: int = 3) -> list:
    """
    Attempt to fetch Indeed jobs via their hidden JSON API endpoints.
    Indeed serves JSON data when specific headers/params are used.
    """
    domain = country.get("indeed_domain", "www.indeed.com")
    kw_enc = urllib.parse.quote_plus(keyword)
    loc_enc = urllib.parse.quote_plus(country["name"])

    jobs = []
    headers = {
        "User-Agent": random_ua(),
        "Accept": "application/json, text/html, */*",
        "Accept-Language": "en-US,en;q=0.9",
        "Referer": f"https://{domain}/jobs?q={kw_enc}",
    }

    for page in range(max_pages):
        start = page * 10
        # Indeed's /jobs endpoint sometimes returns JSON with the right Accept header
        url = f"https://{domain}/jobs?q={kw_enc}&l={loc_enc}&fromage={days}&sort=date&start={start}"

        try:
            resp = requests.get(url, headers=headers, timeout=15)
            if resp.status_code != 200:
                break

            text = resp.text

            # Try to extract job data from the HTML/JSON hybrid response
            # Indeed embeds job data in window._initialData or mosaic-provider-jobcards
            json_match = re.search(r'window\._initialData\s*=\s*(\{.*?\});\s*</script>', text, re.DOTALL)
            if not json_match:
                json_match = re.search(r'"mosaic-provider-jobcards":\s*(\{.*?\})\s*[,}]', text, re.DOTALL)

            if json_match:
                try:
                    data = json.loads(json_match.group(1))
                    # Navigate the nested JSON structure
                    results = (
                        data.get("jobResults", {}).get("results", []) or
                        data.get("results", []) or
                        data.get("searchResults", {}).get("results", []) or
                        []
                    )
                    for result in results:
                        title = result.get("title", "") or result.get("jobTitle", "")
                        company = result.get("company", "") or result.get("companyName", "")
                        location = result.get("formattedLocation", "") or result.get("location", "")
                        jk = result.get("jobkey", "") or result.get("jk", "")
                        posted = result.get("formattedRelativeTime", "") or result.get("pubDate", "")

                        if not title or not jk:
                            continue

                        job_url = f"https://www.indeed.com/viewjob?jk={jk}"
                        desc = result.get("snippet", "") or result.get("description", "")

                        jobs.append(_make_job_record(
                            title=title, company=company,
                            location=location or country["name"],
                            country=country, keyword=keyword, posted_date=posted,
                            source="Indeed-API", job_url=job_url, desc=desc,
                            easy_apply="Yes" if result.get("indeedApply") else "No",
                        ))
                except (json.JSONDecodeError, KeyError):
                    pass

            if not json_match:
                # Fallback: extract job keys from HTML using regex patterns
                jk_matches = re.findall(r'data-jk="([a-zA-Z0-9]+)"', text)
                title_matches = re.findall(r'<h2[^>]*jobTitle[^>]*>.*?<span[^>]*>([^<]+)</span>', text, re.DOTALL)
                company_matches = re.findall(r'data-testid="company-name"[^>]*>([^<]+)<', text)

                for idx, jk in enumerate(jk_matches):
                    title = title_matches[idx].strip() if idx < len(title_matches) else keyword
                    company = company_matches[idx].strip() if idx < len(company_matches) else "Company not stated"

                    jobs.append(_make_job_record(
                        title=title, company=company,
                        location=country["name"], country=country,
                        keyword=keyword, posted_date=datetime.now().strftime("%Y-%m-%d"),
                        source="Indeed-API", job_url=f"https://www.indeed.com/viewjob?jk={jk}",
                    ))

            delay(0.8, 1.5)
        except Exception as e:
            log(f"    Indeed JSON API warning ({country['name']} page {page}): {e}", "WARN")
            break

    return jobs


# ══════════════════════════════════════════════════════════════════════
# TIER 1B — GLASSDOOR GRAPHQL API (Key fix for Glassdoor scraping)
# ══════════════════════════════════════════════════════════════════════

def _get_glassdoor_csrf_token() -> str:
    """
    Extract Glassdoor's CSRF token required for GraphQL API calls.
    The token is set as a cookie when loading any Glassdoor page.
    Caches to disk to avoid repeated requests.
    """
    # Check cache first
    if GD_CSRF_CACHE.exists():
        try:
            cache = json.loads(GD_CSRF_CACHE.read_text())
            token = cache.get("token", "")
            expires = cache.get("expires", 0)
            if token and time.time() < expires:
                return token
        except Exception:
            pass

    # Fetch a lightweight page to get the CSRF cookie
    headers = {
        "User-Agent": random_ua(),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
    }

    token = ""
    try:
        session = requests.Session()
        resp = session.get("https://www.glassdoor.com/Job/jobs.htm", headers=headers,
                          timeout=15, allow_redirects=True)

        # Extract CSRF from cookies
        for cookie in session.cookies:
            if cookie.name in ("gdToken", "gd-csrf-token", "GSESSIONID"):
                token = cookie.value
                break

        # Also try extracting from response headers or meta tags
        if not token:
            csrf_header = resp.headers.get("gd-csrf-token", "")
            if csrf_header:
                token = csrf_header

        if not token:
            # Extract from HTML meta tag or inline script
            m = re.search(r'"gdToken"\s*:\s*"([^"]+)"', resp.text)
            if m:
                token = m.group(1)
            else:
                m = re.search(r'csrfToken\s*[=:]\s*["\']([^"\']+)["\']', resp.text)
                if m:
                    token = m.group(1)

        if not token:
            # Generate a token format Glassdoor accepts for unauthenticated requests
            token = hashlib.md5(f"glassdoor-{int(time.time())}".encode()).hexdigest()

        # Cache the token (valid for 1 hour)
        SESSION_DIR.mkdir(parents=True, exist_ok=True)
        GD_CSRF_CACHE.write_text(json.dumps({
            "token": token,
            "expires": time.time() + 3600,
        }))

    except Exception as e:
        log(f"    Glassdoor CSRF extraction warning: {e}", "WARN")
        token = hashlib.md5(f"glassdoor-{int(time.time())}".encode()).hexdigest()

    return token


def fetch_glassdoor_graphql(country: dict, keyword: str, days: int = 2, max_pages: int = 3) -> list:
    """
    Fetch Glassdoor jobs via their internal GraphQL API.
    This mirrors the LinkedIn Voyager approach — hitting the structured JSON
    endpoint that the frontend uses, rather than scraping DOM elements.
    """
    csrf_token = _get_glassdoor_csrf_token()
    jobs = []

    gd_loc_id = country.get("gd_location_id", 0)
    gd_loc_type = country.get("gd_location_type", "N")

    session = requests.Session()
    session.headers.update({
        "User-Agent": random_ua(),
        "Accept": "*/*",
        "Accept-Language": "en-US,en;q=0.9",
        "Content-Type": "application/json",
        "gd-csrf-token": csrf_token,
        "apollographql-client-name": "job-search-next",
        "apollographql-client-version": "5.0.0",
        "Origin": "https://www.glassdoor.com",
        "Referer": "https://www.glassdoor.com/Job/jobs.htm",
    })
    session.cookies.set("gdId", hashlib.md5(str(time.time()).encode()).hexdigest()[:32],
                        domain=".glassdoor.com")

    for page_num in range(1, max_pages + 1):
        # GraphQL query matching Glassdoor's internal frontend requests
        graphql_payload = {
            "operationName": "JobSearchResultsQuery",
            "variables": {
                "keyword": keyword,
                "locationId": gd_loc_id,
                "locationType": gd_loc_type,
                "numPerPage": 30,
                "pageNumber": page_num,
                "filterParams": [
                    {"filterKey": "fromAge", "values": [str(days)]},
                    {"filterKey": "sortBy", "values": ["date_desc"]},
                ],
                "parameterUrlInput": f"KO0,{len(keyword)}",
                "seoUrl": False,
            },
            "query": """
            query JobSearchResultsQuery(
                $keyword: String!
                $locationId: Int
                $locationType: String
                $numPerPage: Int
                $pageNumber: Int
                $filterParams: [FilterParams]
                $parameterUrlInput: String
                $seoUrl: Boolean
            ) {
                jobListings(
                    contextHolder: {
                        searchParams: {
                            keyword: $keyword
                            locationId: $locationId
                            locationType: $locationType
                            numPerPage: $numPerPage
                            pageNumber: $pageNumber
                            filterParams: $filterParams
                            parameterUrlInput: $parameterUrlInput
                            seoUrl: $seoUrl
                        }
                    }
                ) {
                    compactJobListings {
                        jobListings {
                            jobview {
                                job {
                                    jobTitleText
                                    listingId
                                    descriptionFragment
                                    jobSource
                                    discoverDate
                                }
                                header {
                                    employerNameFromSearch
                                    locationName
                                    payPeriod
                                    payPeriodAdjustedPay {
                                        p50
                                    }
                                    sponsored
                                    easyApply
                                    urgencySignal {
                                        labelText
                                    }
                                    ageInDays
                                    goc
                                    adOrderId
                                    adOrderSponsorshipLevel
                                    jobLink
                                    organic
                                }
                            }
                        }
                        totalJobsCount
                        jobsPageSeoData {
                            pageMetaTitle
                        }
                    }
                }
            }
            """,
        }

        try:
            resp = session.post("https://www.glassdoor.com/graph",
                               json=graphql_payload, timeout=20)

            if resp.status_code == 200:
                try:
                    data = resp.json()
                except (json.JSONDecodeError, ValueError):
                    log(f"    Glassdoor GraphQL: non-JSON response on page {page_num}", "WARN")
                    break

                # Navigate the GraphQL response structure
                listings = []
                try:
                    compact = data.get("data", {}).get("jobListings", {}).get("compactJobListings", {})
                    listings = compact.get("jobListings", [])
                    total_count = compact.get("totalJobsCount", 0)
                    if page_num == 1:
                        log(f"    Glassdoor GraphQL: {total_count} total jobs for '{keyword}' in {country['name']}", "INFO")
                except (AttributeError, TypeError):
                    pass

                if not listings:
                    # Try alternative response shapes
                    listings = (
                        data.get("data", {}).get("jobListings", {}).get("jobListings", []) or
                        data.get("jobListings", []) or
                        []
                    )

                if not listings:
                    break

                for listing in listings:
                    try:
                        jobview = listing.get("jobview", listing)
                        job_data = jobview.get("job", {})
                        header = jobview.get("header", {})

                        title = job_data.get("jobTitleText", "") or ""
                        listing_id = str(job_data.get("listingId", ""))
                        desc_fragment = job_data.get("descriptionFragment", "") or ""
                        company = header.get("employerNameFromSearch", "") or ""
                        location = header.get("locationName", "") or country["name"]
                        easy_apply = "Yes" if header.get("easyApply") else "No"
                        age_in_days = header.get("ageInDays")
                        job_link = header.get("jobLink", "")

                        if not title or not listing_id:
                            continue

                        # Build Glassdoor job URL
                        if job_link and job_link.startswith("/"):
                            job_url = f"https://www.glassdoor.com{job_link}"
                        elif job_link:
                            job_url = job_link
                        else:
                            job_url = f"https://www.glassdoor.com/job-listing/?jl={listing_id}"

                        # Calculate posted date from age
                        if age_in_days is not None:
                            posted = (datetime.now() - timedelta(days=int(age_in_days))).strftime("%Y-%m-%d")
                        else:
                            discover = job_data.get("discoverDate")
                            posted = normalise_posted_date(discover) if discover else datetime.now().strftime("%Y-%m-%d")

                        jobs.append(_make_job_record(
                            title=title, company=company, location=location,
                            country=country, keyword=keyword, posted_date=posted,
                            source="Glassdoor-API", job_url=job_url, desc=desc_fragment,
                            easy_apply=easy_apply,
                        ))
                    except Exception:
                        continue

                if len(listings) < 10:
                    break  # Last page

            elif resp.status_code == 403:
                log(f"    Glassdoor GraphQL: 403 Forbidden (CSRF token may be expired)", "WARN")
                # Invalidate cache and retry with fresh token
                if GD_CSRF_CACHE.exists():
                    GD_CSRF_CACHE.unlink()
                break
            elif resp.status_code == 429:
                log(f"    Glassdoor GraphQL: rate limited (429) — sleeping 5s", "WARN")
                time.sleep(5)
            else:
                log(f"    Glassdoor GraphQL: HTTP {resp.status_code} on page {page_num}", "WARN")
                break

            delay(1.5, 3.0)

        except Exception as e:
            log(f"    Glassdoor GraphQL error ({country['name']} page {page_num}): {e}", "WARN")
            break

    return jobs


def fetch_glassdoor_web_api(country: dict, keyword: str, days: int = 2) -> list:
    """
    Alternative Glassdoor endpoint: scrape search results page HTML
    and extract structured job data from embedded JSON (<script> tags).
    Similar to how LinkedIn embeds Voyager data in <code> tags.
    """
    kw_enc = urllib.parse.quote_plus(keyword)
    loc_enc = urllib.parse.quote_plus(country["name"])

    url = f"https://www.glassdoor.com/Job/jobs.htm?sc.keyword={kw_enc}&locT=N&locKeyword={loc_enc}&fromAge={days}"
    headers = {
        "User-Agent": random_ua(),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
    }

    jobs = []
    try:
        resp = requests.get(url, headers=headers, timeout=15, allow_redirects=True)
        if resp.status_code != 200:
            return jobs

        # Extract embedded JSON data from the page
        # Glassdoor embeds __NEXT_DATA__ or apolloState in script tags
        patterns = [
            r'<script id="__NEXT_DATA__"[^>]*>\s*(\{.*?\})\s*</script>',
            r'window\.__APOLLO_STATE__\s*=\s*(\{.*?\});\s*</script>',
            r'"jobListings"\s*:\s*(\[.*?\])\s*[,}]',
            r'"jobview"\s*:\s*(\{.*?\})\s*[,}]',
        ]

        for pattern in patterns:
            match = re.search(pattern, resp.text, re.DOTALL)
            if not match:
                continue

            try:
                raw_json = match.group(1)
                data = json.loads(raw_json)

                # Navigate various possible structures
                listings = []
                if isinstance(data, list):
                    listings = data
                elif isinstance(data, dict):
                    # __NEXT_DATA__ format
                    props = data.get("props", {}).get("pageProps", {})
                    listings = (
                        props.get("jobListings", {}).get("jobListings", []) or
                        props.get("jobs", []) or
                        []
                    )

                    # Apollo state format — extract job entities
                    if not listings:
                        for key, val in data.items():
                            if isinstance(val, dict) and val.get("__typename") == "JobView":
                                listings.append({"jobview": val})

                for listing in listings:
                    jobview = listing.get("jobview", listing)
                    job = jobview.get("job", jobview)
                    header = jobview.get("header", jobview)

                    title = job.get("jobTitleText", "") or job.get("title", "")
                    listing_id = str(job.get("listingId", "") or job.get("id", ""))
                    company = header.get("employerNameFromSearch", "") or header.get("company", "")
                    location = header.get("locationName", "") or header.get("location", "")
                    age = header.get("ageInDays")

                    if not title:
                        continue

                    job_url = f"https://www.glassdoor.com/job-listing/?jl={listing_id}" if listing_id else ""

                    if age is not None:
                        posted = (datetime.now() - timedelta(days=int(age))).strftime("%Y-%m-%d")
                    else:
                        posted = datetime.now().strftime("%Y-%m-%d")

                    jobs.append(_make_job_record(
                        title=title, company=company, location=location or country["name"],
                        country=country, keyword=keyword, posted_date=posted,
                        source="Glassdoor-Web", job_url=job_url,
                        desc=job.get("descriptionFragment", ""),
                    ))

                if jobs:
                    break  # Found data, stop trying patterns

            except (json.JSONDecodeError, KeyError, TypeError):
                continue

    except Exception as e:
        log(f"    Glassdoor web API warning ({country['name']}): {e}", "WARN")

    return jobs


# ══════════════════════════════════════════════════════════════════════
# TIER 2 — GOOGLE SEARCH FALLBACK
# ══════════════════════════════════════════════════════════════════════

def _google_search_jobs(site_domain: str, keyword: str, country_name: str,
                        source_label: str, country: dict, max_results: int = 30) -> list:
    """
    Use Google search to find job listings on a specific site.
    Parses Google search results HTML to extract job URLs and metadata.
    """
    jobs = []
    kw_enc = urllib.parse.quote_plus(f'site:{site_domain} "{keyword}" "{country_name}"')

    headers = {
        "User-Agent": random_ua(),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate",
    }

    seen_urls = set()
    for start in range(0, max_results, 10):
        url = f"https://www.google.com/search?q={kw_enc}&start={start}&num=10"

        try:
            resp = requests.get(url, headers=headers, timeout=15)
            if resp.status_code == 429:
                log(f"    Google rate limited — stopping search for {source_label}", "WARN")
                break
            if resp.status_code != 200:
                break

            text = resp.text

            # Extract URLs from Google results
            url_patterns = [
                # Standard search result links
                rf'href="(https?://(?:www\.)?{re.escape(site_domain)}[^"]*(?:viewjob|job-listing|jobs)[^"]*)"',
                # Redirect URLs
                rf'/url\?q=(https?://(?:www\.)?{re.escape(site_domain)}[^&]*)',
            ]

            for pattern in url_patterns:
                for match in re.finditer(pattern, text):
                    raw_url = urllib.parse.unquote(match.group(1))
                    norm = normalize_url(raw_url)
                    if norm and norm not in seen_urls:
                        seen_urls.add(norm)

            # Extract titles from search results — correlate with URLs
            title_blocks = re.findall(
                r'<h3[^>]*>(.*?)</h3>.*?<cite[^>]*>(.*?)</cite>',
                text, re.DOTALL
            )

            result_blocks = re.findall(
                r'<div class="[^"]*"[^>]*>.*?<a href="([^"]+)"[^>]*>.*?<h3[^>]*>(.*?)</h3>.*?<span[^>]*>(.*?)</span>',
                text, re.DOTALL
            )

            for block_url, block_title, block_snippet in result_blocks:
                clean_url = urllib.parse.unquote(block_url)
                if site_domain not in clean_url:
                    continue
                norm = normalize_url(clean_url)
                if not norm or norm in seen_urls:
                    continue
                seen_urls.add(norm)

                title = re.sub(r'<[^>]+>', '', block_title).strip()
                snippet = re.sub(r'<[^>]+>', '', block_snippet).strip()

                # Parse title: "DevOps Engineer - Company Name | Glassdoor"
                title_parts = re.split(r'\s*[-|]\s*', title)
                job_title = title_parts[0].strip() if title_parts else keyword
                company = title_parts[1].strip() if len(title_parts) > 1 else "Company not stated"
                # Remove "Glassdoor" / "Indeed" from company name
                company = re.sub(r'\b(Glassdoor|Indeed|Jobs?)\b', '', company, flags=re.IGNORECASE).strip()
                if not company or len(company) < 2:
                    company = "Company not stated"

                jobs.append(_make_job_record(
                    title=job_title, company=company, location=country_name,
                    country=country, keyword=keyword,
                    posted_date=datetime.now().strftime("%Y-%m-%d"),
                    source=source_label, job_url=norm, desc=snippet,
                ))

            delay(2.0, 4.0)  # Respect Google rate limits

        except Exception as e:
            log(f"    Google search error ({source_label}): {e}", "WARN")
            break

    # If we found URLs but couldn't parse titles, create basic records
    if not jobs and seen_urls:
        for url in list(seen_urls)[:max_results]:
            jobs.append(_make_job_record(
                title=keyword, company="See listing",
                location=country_name, country=country, keyword=keyword,
                posted_date=datetime.now().strftime("%Y-%m-%d"),
                source=source_label, job_url=url,
            ))

    return jobs


def fetch_indeed_google(country: dict, keyword: str) -> list:
    """Fetch Indeed jobs via Google search fallback."""
    domain = country.get("indeed_domain", "www.indeed.com")
    return _google_search_jobs(domain, keyword, country["name"],
                               "Indeed-Google", country)


def fetch_glassdoor_google(country: dict, keyword: str) -> list:
    """Fetch Glassdoor jobs via Google search fallback."""
    return _google_search_jobs("glassdoor.com", keyword, country["name"],
                               "Glassdoor-Google", country)


# ══════════════════════════════════════════════════════════════════════
# TIER 3 — BROWSER WITH XHR INTERCEPTION (Playwright)
# ══════════════════════════════════════════════════════════════════════

def crawl_with_xhr_interception(page, country: dict, keyword: str,
                                source: str = "indeed", days: int = 2,
                                max_jobs: int = 50) -> list:
    """
    Use Playwright to navigate to Indeed/Glassdoor and intercept the XHR/fetch
    JSON responses instead of scraping DOM elements.
    This is immune to CSS selector changes — we capture the same data
    the frontend renders.
    """
    jobs = []
    captured_responses = []

    def on_response(response):
        """Capture JSON responses from AJAX calls."""
        try:
            url = response.url
            content_type = response.headers.get("content-type", "")

            # Indeed AJAX patterns
            if source == "indeed" and ("indeed.com" in url):
                if ("application/json" in content_type or
                    "json" in content_type or
                    "/rpc/" in url or
                    "graphql" in url or
                    "mosaic" in url or
                    "jobResults" in url or
                    "api" in url):
                    try:
                        body = response.json()
                        captured_responses.append({"url": url, "data": body, "source": "indeed"})
                    except Exception:
                        pass

            # Glassdoor AJAX/GraphQL patterns
            elif source == "glassdoor" and ("glassdoor.com" in url):
                if ("application/json" in content_type or
                    "json" in content_type or
                    "/graph" in url or
                    "api" in url):
                    try:
                        body = response.json()
                        captured_responses.append({"url": url, "data": body, "source": "glassdoor"})
                    except Exception:
                        pass
        except Exception:
            pass

    # Register the response interceptor
    page.on("response", on_response)

    try:
        if source == "indeed":
            domain = country.get("indeed_domain", "www.indeed.com")
            kw_enc = urllib.parse.quote_plus(keyword)
            loc_enc = urllib.parse.quote_plus(country["name"])
            url = f"https://{domain}/jobs?q={kw_enc}&l={loc_enc}&fromage={days}&sort=date"
        else:  # glassdoor
            kw_enc = urllib.parse.quote_plus(keyword)
            gd_loc = country.get("gd_location_id", 0)
            url = f"https://www.glassdoor.com/Job/jobs.htm?sc.keyword={kw_enc}&locId={gd_loc}&locT=N&fromAge={days}"

        try:
            page.goto(url, wait_until="domcontentloaded", timeout=15_000)
        except Exception:
            pass

        # Explicitly wait for job cards to render in the DOM
        try:
            if source == "indeed":
                page.wait_for_selector("div.job_seen_beacon, td.resultContent, a[data-jk]", timeout=8_000)
            else:
                page.wait_for_selector("a[data-test='job-title'], li[data-test='jobListing'], a.JobCard_jobTitle__GLyJ1", timeout=8_000)
        except Exception:
            pass

        delay(1.5, 3.0)

        # Scroll to trigger lazy loading and additional AJAX calls
        for _ in range(2):
            try:
                page.keyboard.press("End")
            except Exception:
                pass
            delay(0.8, 1.2)

        # 1. First Priority: DOM scraping (extracts full job titles, companies, locations, URLs)
        if source == "indeed":
            jobs.extend(_dom_scrape_indeed(page, country, keyword))
        else:
            jobs.extend(_dom_scrape_glassdoor(page, country, keyword))

        if jobs:
            log(f"    DOM scraping: {len(jobs)} jobs from {source}", "OK")
            return jobs[:max_jobs]

        # 2. Second Priority: Process captured JSON responses
        for captured in captured_responses:
            data = captured["data"]
            if captured["source"] == "indeed":
                jobs.extend(_parse_indeed_xhr(data, country, keyword))
            elif captured["source"] == "glassdoor":
                jobs.extend(_parse_glassdoor_xhr(data, country, keyword))

        if jobs:
            log(f"    XHR interception: {len(jobs)} jobs from {source}", "OK")
            return jobs[:max_jobs]

        # 3. Third Priority: Embedded JSON extraction
        try:
            page_source = page.content()
            if source == "indeed":
                jobs.extend(_extract_indeed_embedded_json(page_source, country, keyword))
            else:
                jobs.extend(_extract_glassdoor_embedded_json(page_source, country, keyword))
        except Exception:
            pass

        if jobs:
            log(f"    Embedded JSON extraction: {len(jobs)} jobs from {source}", "OK")
            return jobs[:max_jobs]

    except Exception as e:
        log(f"    Browser {source} error ({country['name']}): {e}", "WARN")
    finally:
        try:
            page.remove_listener("response", on_response)
        except Exception:
            pass

    return jobs[:max_jobs]


def _parse_indeed_xhr(data, country: dict, keyword: str) -> list:
    """Parse Indeed XHR/AJAX JSON response into job records."""
    jobs = []
    if not isinstance(data, dict):
        if isinstance(data, list):
            results = data
        else:
            return jobs
    else:
        results = (
            data.get("jobResults", {}).get("results", []) or
            data.get("results", []) or
            data.get("searchResults", {}).get("results", []) or
            data.get("metaData", {}).get("mosaicProviderJobCardsModel", {}).get("results", []) or
            []
        )
        if not results:
            mosaic = data.get("mosaic", {}).get("providerData", {})
            if isinstance(mosaic, dict):
                for key, val in mosaic.items():
                    if isinstance(val, dict) and "results" in val:
                        results = val["results"]
                        break

    if not isinstance(results, list):
        return jobs

    for result in results:
        if not isinstance(result, dict):
            continue
        title = result.get("title", "") or result.get("jobTitle", "")
        company = result.get("company", "") or result.get("companyName", "")
        location = result.get("formattedLocation", "") or result.get("location", "")
        jk = result.get("jobkey", "") or result.get("jk", "")
        posted = result.get("formattedRelativeTime", "") or result.get("pubDate", "")
        desc = result.get("snippet", "") or result.get("description", "")

        if not title or not jk:
            continue

        jobs.append(_make_job_record(
            title=title, company=company,
            location=location or country["name"],
            country=country, keyword=keyword, posted_date=posted,
            source="Indeed-Browser", job_url=f"https://www.indeed.com/viewjob?jk={jk}",
            desc=desc,
            easy_apply="Yes" if result.get("indeedApply") else "No",
        ))

    return jobs


def _parse_glassdoor_xhr(data, country: dict, keyword: str) -> list:
    """Parse Glassdoor GraphQL/AJAX JSON response into job records."""
    jobs = []
    if not isinstance(data, dict):
        if isinstance(data, list):
            listings = data
        else:
            return jobs
    else:
        listings = []
        try:
            compact = data.get("data", {}).get("jobListings", {}).get("compactJobListings", {})
            if isinstance(compact, dict):
                listings = compact.get("jobListings", [])
        except (AttributeError, TypeError):
            pass

        if not listings:
            listings = (
                data.get("data", {}).get("jobListings", {}).get("jobListings", []) or
                data.get("jobListings", []) or
                []
            )

    if not isinstance(listings, list):
        return jobs

    for listing in listings:
        try:
            if not isinstance(listing, dict):
                continue
            jobview = listing.get("jobview", listing)
            if not isinstance(jobview, dict):
                continue
            job_data = jobview.get("job", {})
            header = jobview.get("header", {})

            title = job_data.get("jobTitleText", "") or ""
            listing_id = str(job_data.get("listingId", ""))
            company = header.get("employerNameFromSearch", "") or ""
            location = header.get("locationName", "") or country["name"]
            age = header.get("ageInDays")

            if not title:
                continue

            if age is not None:
                posted = (datetime.now() - timedelta(days=int(age))).strftime("%Y-%m-%d")
            else:
                posted = datetime.now().strftime("%Y-%m-%d")

            job_url = f"https://www.glassdoor.com/job-listing/?jl={listing_id}" if listing_id else ""

            jobs.append(_make_job_record(
                title=title, company=company, location=location,
                country=country, keyword=keyword, posted_date=posted,
                source="Glassdoor-Browser", job_url=job_url,
                desc=job_data.get("descriptionFragment", ""),
                easy_apply="Yes" if header.get("easyApply") else "No",
            ))
        except Exception:
            continue

    return jobs


def _extract_indeed_embedded_json(html: str, country: dict, keyword: str) -> list:
    """Extract job data from Indeed's embedded JSON in page source."""
    jobs = []
    patterns = [
        r'window\._initialData\s*=\s*(\{.*?\});\s*</script>',
        r'window\.mosaic\.providerData\["mosaic-provider-jobcards"\]\s*=\s*(\{.*?\});\s*</script>',
    ]

    for pattern in patterns:
        match = re.search(pattern, html, re.DOTALL)
        if match:
            try:
                data = json.loads(match.group(1))
                return _parse_indeed_xhr(data, country, keyword)
            except (json.JSONDecodeError, KeyError):
                continue

    # Regex fallback for job keys
    jk_matches = re.findall(r'data-jk="([a-zA-Z0-9]+)"', html)
    for jk in dict.fromkeys(jk_matches):
        jobs.append(_make_job_record(
            title=keyword, company="See listing",
            location=country["name"], country=country, keyword=keyword,
            posted_date=datetime.now().strftime("%Y-%m-%d"),
            source="Indeed-Browser", job_url=f"https://www.indeed.com/viewjob?jk={jk}",
        ))

    return jobs


def _extract_glassdoor_embedded_json(html: str, country: dict, keyword: str) -> list:
    """Extract job data from Glassdoor's embedded JSON (__NEXT_DATA__)."""
    jobs = []
    patterns = [
        r'<script id="__NEXT_DATA__"[^>]*>\s*(\{.*?\})\s*</script>',
        r'window\.__APOLLO_STATE__\s*=\s*(\{.*?\});\s*</script>',
    ]

    for pattern in patterns:
        match = re.search(pattern, html, re.DOTALL)
        if match:
            try:
                data = json.loads(match.group(1))
                # __NEXT_DATA__ structure
                props = data.get("props", {}).get("pageProps", {})
                listings = (
                    props.get("jobListings", {}).get("jobListings", []) or
                    props.get("jobs", []) or
                    []
                )
                for listing in listings:
                    jobview = listing.get("jobview", listing)
                    job = jobview.get("job", jobview)
                    header = jobview.get("header", jobview)

                    title = job.get("jobTitleText", "") or job.get("title", "")
                    listing_id = str(job.get("listingId", "") or job.get("id", ""))
                    company = header.get("employerNameFromSearch", "") or header.get("company", "")
                    location = header.get("locationName", "") or country["name"]

                    if not title:
                        continue

                    age = header.get("ageInDays")
                    if age is not None:
                        posted = (datetime.now() - timedelta(days=int(age))).strftime("%Y-%m-%d")
                    else:
                        posted = datetime.now().strftime("%Y-%m-%d")

                    jobs.append(_make_job_record(
                        title=title, company=company, location=location,
                        country=country, keyword=keyword, posted_date=posted,
                        source="Glassdoor-Browser", 
                        job_url=f"https://www.glassdoor.com/job-listing/?jl={listing_id}",
                        desc=job.get("descriptionFragment", ""),
                    ))

                if jobs:
                    return jobs

            except (json.JSONDecodeError, KeyError, TypeError):
                continue

    return jobs


def _dom_scrape_indeed(page, country: dict, keyword: str) -> list:
    """DOM scrape Indeed search results."""
    jobs = []
    try:
        cards = page.query_selector_all("div.job_seen_beacon, td.resultContent, div.cardOutline, li div[class*='card']")
        for card in cards[:50]:
            try:
                title_el = card.query_selector("a[data-jk] span, [class*='jobTitle'] span, h2 span, a.jcs-JobTitle span, [data-testid='jobTitle']")
                if not title_el:
                    title_el = card.query_selector("a[data-jk], a.jcs-JobTitle, h2")
                
                comp_el = card.query_selector("[data-testid='company-name'], span.companyName, [class*='companyName'], span[class*='company']")
                loc_el = card.query_selector("[data-testid='text-location'], div.companyLocation, [class*='companyLocation'], div[class*='location']")
                link_el = card.query_selector("a[data-jk], a.jcs-JobTitle")

                title = title_el.inner_text().strip() if title_el else ""
                company = comp_el.inner_text().strip() if comp_el else "Company not stated"
                location = loc_el.inner_text().strip() if loc_el else country["name"]

                href = ""
                domain = country.get("indeed_domain", "www.indeed.com")
                if link_el:
                    href = link_el.get_attribute("href") or ""
                    if href.startswith("/"):
                        href = f"https://{domain}{href}"

                jk = card.get_attribute("data-jk") or (link_el.get_attribute("data-jk") if link_el else "")
                if jk:
                    href = f"https://www.indeed.com/viewjob?jk={jk}"

                if not title or not href:
                    continue

                jobs.append(_make_job_record(
                    title=title, company=company, location=location,
                    country=country, keyword=keyword,
                    posted_date=datetime.now().strftime("%Y-%m-%d"),
                    source="Indeed-Browser", job_url=href,
                ))
            except Exception:
                continue
    except Exception:
        pass
    return jobs


def _dom_scrape_glassdoor(page, country: dict, keyword: str) -> list:
    """DOM scrape Glassdoor search results."""
    jobs = []
    try:
        cards = page.query_selector_all(
            "li[data-test='jobListing'], "
            "div[data-test='job-card-wrapper'], "
            "li[class*='jobListingCard'], "
            "div[class*='JobCard_jobCard'], "
            "div.jobCard, "
            "li.react-job-listing"
        )
        for card in cards[:50]:
            try:
                title_el = card.query_selector("a[data-test='job-title'], a.JobCard_jobTitle__GLyJ1, a[class*='jobTitle'], a.job-title")
                comp_el = card.query_selector(
                    "span[class*='EmployerName'], "
                    "div[class*='EmployerProfile'], "
                    "span[class*='compactEmployerName'], "
                    "[data-test='emp-name']"
                )
                loc_el = card.query_selector(
                    "div[data-test='emp-location'], "
                    "span[data-test='emp-location'], "
                    "div[class*='location'], "
                    "[class*='location']"
                )

                title = title_el.inner_text().strip() if title_el else ""
                company_raw = comp_el.inner_text().strip() if comp_el else "Company not stated"
                # Strip star ratings (e.g. "Leidos\n3.8")
                company = re.sub(r'[\r\n\s]+\d+\.\d+$', '', company_raw).strip()
                location = loc_el.inner_text().strip() if loc_el else country["name"]

                href = ""
                if title_el:
                    href = title_el.get_attribute("href") or ""
                    if href.startswith("/"):
                        href = f"https://www.glassdoor.com{href}"

                jl = card.get_attribute("data-job-id") or card.get_attribute("data-id") or ""
                if not href and jl:
                    href = f"https://www.glassdoor.com/job-listing/?jl={jl}"

                if not title or not href:
                    continue

                jobs.append(_make_job_record(
                    title=title, company=company, location=location,
                    country=country, keyword=keyword,
                    posted_date=datetime.now().strftime("%Y-%m-%d"),
                    source="Glassdoor-Browser", job_url=href,
                ))
            except Exception:
                continue
    except Exception:
        pass
    return jobs


# ══════════════════════════════════════════════════════════════════════
# WORD & EXCEL GENERATORS (preserved from original)
# ══════════════════════════════════════════════════════════════════════
def save_multi_sheet_excel(df: pd.DataFrame):
    """Save clean multi-sheet Excel with 'All Jobs' master sheet + Date Tabs."""
    if not HAS_EXCEL or df is None or df.empty:
        return

    clean_df = deduplicate_dataframe(df)
    log("\nGenerating Indeed & Glassdoor multi-sheet Excel (Master + Date Tabs) …")
    wb = openpyxl.Workbook()
    if "Sheet" in wb.sheetnames:
        wb.remove(wb["Sheet"])

    cols = ["#","Crawl Date","Source","Country","Flag","Tier","City","City Openings",
            "Location","Company","Job Title","Search Keyword","Posted Date",
            "Easy Apply","Remote / Workplace","Match Score","Visa Sponsorship",
            "Skills","Resume","Applied","Notes","Job URL"]

    hf   = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    hfnt = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
    thin = Border(left=Side(style="thin"), right=Side(style="thin"),
                  top=Side(style="thin"),  bottom=Side(style="thin"))
    vf   = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")

    def populate_sheet(ws, sub_df):
        for ci, col in enumerate(cols, 1):
            c = ws.cell(row=1, column=ci, value=col)
            c.fill = hf; c.font = hfnt
            c.alignment = Alignment(horizontal="center", vertical="center")
            c.border = thin

        for ri, (_, job) in enumerate(sub_df.iterrows(), 2):
            vv = str(job.get("Visa Sponsorship Mentioned", "No"))
            vals = [
                ri-1, str(job.get("Crawl Date", job.get("Date", ""))), str(job.get("Source", "Indeed/Glassdoor")),
                str(job.get("Country","")), str(job.get("Flag","")), str(job.get("Tier","")),
                str(job.get("City","")), str(job.get("City Openings","")),
                str(job.get("Location","")), str(job.get("Company","")),
                str(job.get("Job Title","")), str(job.get("Search Keyword","")),
                str(job.get("Posted Date","")), str(job.get("Easy Apply","No")),
                str(job.get("Remote / Workplace","On-site / Hybrid")),
                str(job.get("Match Score","85%")), vv,
                str(job.get("Required Skills","")), str(job.get("Resume File Path","")),
                str(job.get("Applied Status","No")), str(job.get("Notes","")),
                str(job.get("Job URL",""))
            ]
            for ci, val in enumerate(vals, 1):
                c = ws.cell(row=ri, column=ci, value=val)
                c.font = Font(name="Calibri", size=9); c.border = thin
                if ci == 17 and vv.lower() == "yes":
                    c.fill = vf

        for i, w in enumerate([4,12,12,14,5,5,16,14,20,22,32,20,16,12,18,12,15,30,45,8,20,55], 1):
            ws.column_dimensions[get_column_letter(i)].width = w
        ws.freeze_panes = "A2"
        if len(sub_df) > 0:
            ws.auto_filter.ref = ws.dimensions

    # 1. Master Sheet "All Jobs"
    ws_all = wb.create_sheet(title="All Jobs")
    populate_sheet(ws_all, clean_df)

    # 2. Date-based sheets
    date_col = "Crawl Date" if "Crawl Date" in clean_df.columns else ("Date" if "Date" in clean_df.columns else None)
    if date_col:
        unique_dates = sorted(clean_df[date_col].dropna().astype(str).str.split("T").str[0].unique(), reverse=True)
        for date_str in unique_dates:
            if not date_str or len(date_str) < 5 or date_str == "nan":
                continue
            sub = clean_df[clean_df[date_col].astype(str).str.startswith(date_str)]
            if not sub.empty:
                ws_date = wb.create_sheet(title=date_str[:31])
                populate_sheet(ws_date, sub)

    try:
        MASTER_EXCEL_PATH.parent.mkdir(parents=True, exist_ok=True)
        EXCEL_PATH.parent.mkdir(parents=True, exist_ok=True)
        wb.save(MASTER_EXCEL_PATH)
        wb.save(EXCEL_PATH)
        log(f"Master Excel saved -> {MASTER_EXCEL_PATH} (All Jobs + Date Tabs)", "OK")
        log(f"Run Excel archived -> {EXCEL_PATH}", "OK")
    except PermissionError:
        alt = WORK_DIR / f"indeed_glassdoor_jobs_{int(time.time())}.xlsx"
        wb.save(alt)
        log(f"Excel locked -- saved to {alt}", "WARN")


def save_master_csv(all_jobs: list) -> pd.DataFrame:
    """Save single master CSV with deduplication and metadata preservation."""
    today_str = datetime.now().strftime("%Y-%m-%d")
    incoming_df = pd.DataFrame(all_jobs)

    if MASTER_CSV_PATH.exists():
        try:
            log(f"Loading existing Indeed/Glassdoor master sheet -> {MASTER_CSV_PATH}", "INFO")
            master_df = pd.read_csv(MASTER_CSV_PATH, on_bad_lines="skip")
            combined_df = pd.concat([master_df, incoming_df], ignore_index=True)
        except Exception as e:
            log(f"Could not read existing master CSV ({e}) -- creating new", "WARN")
            combined_df = incoming_df
    else:
        combined_df = incoming_df

    df = deduplicate_dataframe(combined_df)
    df["City"] = df["Location"].apply(extract_city)
    city_counts = df["City"].value_counts().to_dict()
    df["City Openings"] = df["City"].map(city_counts).fillna(0).astype(int)

    RUN_DIR.mkdir(parents=True, exist_ok=True)
    df.to_csv(MASTER_CSV_PATH, index=False)
    df.to_csv(CSV_PATH, index=False)
    log(f"Master CSV updated -> {MASTER_CSV_PATH} ({len(df)} unique jobs maintaining single master sheet)", "OK")
    log(f"Run CSV archived  -> {CSV_PATH}", "OK")

    # Also automatically sync and merge into full_crawl_jobs.csv
    global_master = WORK_DIR / "full_crawl_jobs.csv"
    if global_master.exists():
        try:
            g_df = pd.read_csv(global_master, on_bad_lines="skip")
            combined_all = pd.concat([g_df, incoming_df], ignore_index=True)
            clean_all = deduplicate_dataframe(combined_all)
            clean_all.to_csv(global_master, index=False)
            log(f"Global Master CSV updated -> {global_master} ({len(clean_all)} unified jobs across all platforms)", "OK")
        except Exception as err:
            log(f"Could not merge into {global_master}: {err}", "WARN")

    return df


def generate_word_report(jobs_df: pd.DataFrame):
    """Generate professional Word Document report for Indeed & Glassdoor crawl."""
    if not HAS_DOCX or jobs_df.empty:
        return
    log("\nGenerating Indeed & Glassdoor Word document report …", "INFO")
    doc = Document()

    # Title
    title_p = doc.add_paragraph()
    r = title_p.add_run(f"Indeed & Glassdoor Job Intelligence Report")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(31, 78, 121)

    sub_p = doc.add_paragraph()
    sub_p.add_run(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')} · {len(jobs_df)} Total Roles across {jobs_df['Country'].nunique()} Countries")
    sub_p.runs[0].font.size = Pt(10)
    sub_p.runs[0].font.italic = True

    # Source breakdown
    if "Source" in jobs_df.columns:
        source_counts = jobs_df["Source"].value_counts().to_dict()
        src_p = doc.add_paragraph()
        src_p.add_run("Sources: ")
        src_p.runs[0].bold = True
        for src, count in source_counts.items():
            src_p.add_run(f"{src}: {count}  ")

    # Country summary table
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdrs = ["Country", "Flag", "Roles Found", "Sources"]
    for i, h in enumerate(hdrs):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for country in TARGET_COUNTRIES:
        c_name = country["name"]
        sub = jobs_df[jobs_df["Country"].str.lower() == c_name.lower()]
        if not sub.empty:
            row = table.add_row().cells
            row[0].text = c_name
            row[1].text = country["flag"]
            row[2].text = str(len(sub))
            row[3].text = ", ".join(sub["Source"].unique())

    try:
        DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
        doc.save(DOCX_PATH)
        log(f"Word document saved -> {DOCX_PATH}", "OK")
    except Exception as e:
        log(f"Could not save Word document: {e}", "WARN")


# ══════════════════════════════════════════════════════════════════════
# MAIN CRAWL PIPELINE — 3-TIER API-FIRST ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════

def _crawl_glassdoor_browser_graphql(page, country: dict, keyword: str, days: int = 2) -> list:
    """
    Make Glassdoor GraphQL API calls FROM WITHIN the browser context.
    This bypasses Cloudflare because the browser session already has valid
    cookies and TLS fingerprint. This is the equivalent of the LinkedIn
    Voyager approach — using the browser's authenticated session to call
    the internal API.
    """
    jobs = []

    # JavaScript to execute the GraphQL query using the browser's fetch()
    # which inherits all Cloudflare cookies and session state
    graphql_js = """
    async (args) => {
        const { keyword, locationId, locationType, numPerPage, pageNumber, fromAge } = args;
        try {
            // Get CSRF token from cookies or meta tags
            let csrfToken = '';
            const cookies = document.cookie.split(';');
            for (const c of cookies) {
                const [name, val] = c.trim().split('=');
                if (name === 'gdId' || name === 'gd-csrf-token' || name === 'GSESSIONID') {
                    csrfToken = val;
                    break;
                }
            }
            if (!csrfToken) {
                // Try extracting from page
                const meta = document.querySelector('meta[name="csrf-token"]');
                if (meta) csrfToken = meta.getAttribute('content');
            }
            if (!csrfToken) csrfToken = 'undefined';

            const payload = {
                operationName: "JobSearchResultsQuery",
                variables: {
                    keyword: keyword,
                    locationId: locationId,
                    locationType: locationType,
                    numPerPage: numPerPage,
                    pageNumber: pageNumber,
                    filterParams: [
                        { filterKey: "fromAge", values: [String(fromAge)] },
                        { filterKey: "sortBy", values: ["date_desc"] },
                    ],
                    parameterUrlInput: `KO0,${keyword.length}`,
                    seoUrl: false,
                },
                query: `query JobSearchResultsQuery($keyword: String!, $locationId: Int, $locationType: String, $numPerPage: Int, $pageNumber: Int, $filterParams: [FilterParams], $parameterUrlInput: String, $seoUrl: Boolean) {
                    jobListings(contextHolder: {searchParams: {keyword: $keyword, locationId: $locationId, locationType: $locationType, numPerPage: $numPerPage, pageNumber: $pageNumber, filterParams: $filterParams, parameterUrlInput: $parameterUrlInput, seoUrl: $seoUrl}}) {
                        compactJobListings {
                            jobListings {
                                jobview {
                                    job { jobTitleText listingId descriptionFragment discoverDate }
                                    header { employerNameFromSearch locationName easyApply ageInDays jobLink organic }
                                }
                            }
                            totalJobsCount
                        }
                    }
                }`
            };

            const resp = await fetch('https://www.glassdoor.com/graph', {
                method: 'POST',
                headers: {
                    'Content-Type': 'application/json',
                    'gd-csrf-token': csrfToken,
                    'apollographql-client-name': 'job-search-next',
                    'apollographql-client-version': '5.0.0',
                },
                credentials: 'include',
                body: JSON.stringify(payload),
            });

            if (!resp.ok) {
                return { error: `HTTP ${resp.status}`, status: resp.status };
            }

            const data = await resp.json();
            return data;
        } catch (e) {
            return { error: e.message };
        }
    }
    """

    for page_num in range(1, 4):  # Max 3 pages
        try:
            result = page.evaluate(graphql_js, {
                "keyword": keyword,
                "locationId": country.get("gd_location_id", 0),
                "locationType": country.get("gd_location_type", "N"),
                "numPerPage": 30,
                "pageNumber": page_num,
                "fromAge": days,
            })

            if not result or result.get("error"):
                err = result.get("error", "unknown") if result else "null response"
                log(f"      GD-GraphQL browser error: {err}", "WARN")
                break

            # Parse the GraphQL response
            listings = []
            try:
                compact = result.get("data", {}).get("jobListings", {}).get("compactJobListings", {})
                listings = compact.get("jobListings", [])
                total = compact.get("totalJobsCount", 0)
                if page_num == 1 and total:
                    log(f"      GD-GraphQL via browser: {total} total results", "INFO")
            except (AttributeError, TypeError):
                pass

            if not listings:
                break

            for listing in listings:
                try:
                    jobview = listing.get("jobview", listing)
                    job_data = jobview.get("job", {})
                    header = jobview.get("header", {})

                    title = job_data.get("jobTitleText", "")
                    listing_id = str(job_data.get("listingId", ""))
                    company = header.get("employerNameFromSearch", "")
                    location = header.get("locationName", "") or country["name"]
                    age = header.get("ageInDays")
                    job_link = header.get("jobLink", "")

                    if not title:
                        continue

                    if job_link and job_link.startswith("/"):
                        job_url = f"https://www.glassdoor.com{job_link}"
                    elif job_link:
                        job_url = job_link
                    else:
                        job_url = f"https://www.glassdoor.com/job-listing/?jl={listing_id}" if listing_id else ""

                    if age is not None:
                        posted = (datetime.now() - timedelta(days=int(age))).strftime("%Y-%m-%d")
                    else:
                        posted = datetime.now().strftime("%Y-%m-%d")

                    jobs.append(_make_job_record(
                        title=title, company=company, location=location,
                        country=country, keyword=keyword, posted_date=posted,
                        source="Glassdoor-API", job_url=job_url,
                        desc=job_data.get("descriptionFragment", ""),
                        easy_apply="Yes" if header.get("easyApply") else "No",
                    ))
                except Exception:
                    continue

            if len(listings) < 10:
                break

            delay(1.5, 3.0)

        except Exception as e:
            log(f"      GD-GraphQL browser exception: {e}", "WARN")
            break

    return jobs


def run_crawler(headless: bool, countries_filter, sources: list, max_per_kw: int,
                days: int, skip_google: bool = False, skip_browser: bool = False) -> list:
    """
    Orchestrate the crawl pipeline. Both Indeed and Glassdoor are behind
    Cloudflare, so the browser is the PRIMARY crawl method:

      1. Quick API probe (RSS/GraphQL) — fast but will fail if Cloudflare blocks
      2. Browser-primary crawl — uses Playwright to bypass Cloudflare, then:
         a. Glassdoor: Makes GraphQL calls from within browser context (page.evaluate)
         b. Indeed: Navigates search pages + intercepts XHR + extracts embedded JSON
         c. DOM scraping as absolute last resort
      3. Google Search fallback — if both API and browser fail
    """
    targets = filter_target_countries(TARGET_COUNTRIES, countries_filter)
    all_jobs = []
    tier_stats = {"Indeed-RSS": 0, "Indeed-API": 0, "Glassdoor-API": 0,
                  "Glassdoor-Web": 0, "Indeed-Google": 0, "Glassdoor-Google": 0,
                  "Indeed-Browser": 0, "Glassdoor-Browser": 0,
                  "Indeed-DOM": 0, "Glassdoor-DOM": 0}

    log("=" * 65)
    log("  INDEED & GLASSDOOR JOB CRAWLER — Browser-Primary Pipeline")
    log(f"   Target Markets : {len(targets)} Countries")
    log(f"   Sources        : {', '.join(sources).upper()}")
    log(f"   Time Window    : Past {days} Days")
    log(f"   Headless Mode  : {headless}")
    log("=" * 65)

    # ── STEP 1: Quick API Probe (attempt fast path) ────────────────
    log("\n── Step 1: Quick API Probe ──", "INFO")

    for country in targets:
        log(f"\n{country['flag']} {country['name']}:")

        if "indeed" in sources:
            for kw in SEARCH_TERMS:
                rss_jobs = fetch_indeed_rss(country, kw, days=days)
                if rss_jobs:
                    log(f"   Indeed-RSS | '{kw}' -> {len(rss_jobs)} jobs", "OK")
                    all_jobs.extend(rss_jobs)
                    tier_stats["Indeed-RSS"] += len(rss_jobs)
                delay(0.3, 0.8)

        if "glassdoor" in sources:
            for kw in SEARCH_TERMS[:1]:  # Only probe first keyword to detect Cloudflare
                gql_jobs = fetch_glassdoor_graphql(country, kw, days=days, max_pages=1)
                if gql_jobs:
                    log(f"   GD-GraphQL | '{kw}' -> {len(gql_jobs)} jobs", "OK")
                    all_jobs.extend(gql_jobs)
                    tier_stats["Glassdoor-API"] += len(gql_jobs)
                    # If probe works, fetch remaining keywords
                    for kw2 in SEARCH_TERMS[1:]:
                        gql_jobs2 = fetch_glassdoor_graphql(country, kw2, days=days, max_pages=2)
                        if gql_jobs2:
                            log(f"   GD-GraphQL | '{kw2}' -> {len(gql_jobs2)} jobs", "OK")
                            all_jobs.extend(gql_jobs2)
                            tier_stats["Glassdoor-API"] += len(gql_jobs2)
                        delay(1.0, 2.0)
                break  # Only probe once per country

    indeed_count = sum(v for k, v in tier_stats.items() if "Indeed" in k)
    glassdoor_count = sum(v for k, v in tier_stats.items() if "Glassdoor" in k or "GD" in k)
    log(f"\n   API Probe: Indeed={indeed_count}, Glassdoor={glassdoor_count}")

    # ── STEP 2: Browser-Primary Crawl (main method) ────────────────
    needs_browser_indeed = ("indeed" in sources and indeed_count == 0)
    needs_browser_glassdoor = ("glassdoor" in sources and glassdoor_count == 0)

    if (needs_browser_indeed or needs_browser_glassdoor) and HAS_PLAYWRIGHT and not skip_browser:
        log("\n── Step 2: Browser-Primary Crawl (Stealth Cloudflare Bypass) ──", "INFO")
        SESSION_DIR.mkdir(parents=True, exist_ok=True)

        browser_args = [
            "--disable-blink-features=AutomationControlled",
            "--disable-infobars",
            "--no-sandbox",
            "--disable-dev-shm-usage",
        ]

        try:
            with sync_playwright() as p:
                try:
                    browser = p.chromium.launch_persistent_context(
                        user_data_dir=str(SESSION_DIR),
                        headless=headless,
                        channel="chrome",
                        args=browser_args,
                        ignore_default_args=["--enable-automation"],
                        user_agent=random_ua(),
                        viewport={"width": 1366, "height": 768},
                    )
                except Exception:
                    browser = p.chromium.launch_persistent_context(
                        user_data_dir=str(SESSION_DIR),
                        headless=headless,
                        args=browser_args,
                        ignore_default_args=["--enable-automation"],
                        user_agent=random_ua(),
                        viewport={"width": 1366, "height": 768},
                    )

                page = browser.pages[0] if browser.pages else browser.new_page()
                page.add_init_script("""
                    Object.defineProperty(navigator, 'webdriver', {get: () => undefined});
                    window.navigator.chrome = { runtime: {} };
                """)

                # ── Glassdoor Crawl ──
                if needs_browser_glassdoor:
                    log("\n   ── Glassdoor: Browser Crawl ──")
                    for country in targets:
                        for kw in SEARCH_TERMS:
                            log(f"   [Glassdoor] {country['flag']} {country['name']} | '{kw}' …")
                            gd_jobs = crawl_with_xhr_interception(
                                page, country, kw, source="glassdoor",
                                days=days, max_jobs=max_per_kw)
                            if gd_jobs:
                                log(f"      -> {len(gd_jobs)} jobs", "OK")
                                all_jobs.extend(gd_jobs)
                                for j in gd_jobs:
                                    src = j.get("Source", "Glassdoor-Browser")
                                    tier_stats[src] = tier_stats.get(src, 0) + 1
                            else:
                                log(f"      -> 0 jobs")
                            delay(1.5, 3.0)

                # ── Indeed Crawl ──
                if needs_browser_indeed:
                    log("\n   ── Indeed: Browser Crawl ──")
                    for country in targets:
                        for kw in SEARCH_TERMS:
                            log(f"   [Indeed] {country['flag']} {country['name']} | '{kw}' …")
                            ind_jobs = crawl_with_xhr_interception(
                                page, country, kw, source="indeed",
                                days=days, max_jobs=max_per_kw)
                            if ind_jobs:
                                log(f"      -> {len(ind_jobs)} jobs", "OK")
                                all_jobs.extend(ind_jobs)
                                for j in ind_jobs:
                                    src = j.get("Source", "Indeed-Browser")
                                    tier_stats[src] = tier_stats.get(src, 0) + 1
                            else:
                                log(f"      -> 0 jobs")
                            delay(1.5, 3.0)

                browser.close()

        except Exception as e:
            log(f"Browser crawl exception: {e}", "WARN")

    # ── STEP 3: Google Search Fallback ─────────────────────────────
    if not skip_google:
        indeed_count = sum(v for k, v in tier_stats.items() if "Indeed" in k)
        glassdoor_count = sum(v for k, v in tier_stats.items() if "Glassdoor" in k or "GD" in k)

        needs_google_indeed = ("indeed" in sources and indeed_count == 0)
        needs_google_glassdoor = ("glassdoor" in sources and glassdoor_count == 0)

        if needs_google_indeed or needs_google_glassdoor:
            log("\n── Step 3: Google Search Fallback ──", "INFO")

            for country in targets:
                for kw in SEARCH_TERMS:
                    if needs_google_glassdoor:
                        gd_google_jobs = fetch_glassdoor_google(country, kw)
                        if gd_google_jobs:
                            log(f"   GD-Google  | {country['flag']} '{kw}' -> {len(gd_google_jobs)} jobs", "OK")
                            all_jobs.extend(gd_google_jobs)
                            tier_stats["Glassdoor-Google"] += len(gd_google_jobs)

                    if needs_google_indeed:
                        ind_google_jobs = fetch_indeed_google(country, kw)
                        if ind_google_jobs:
                            log(f"   Ind-Google | {country['flag']} '{kw}' -> {len(ind_google_jobs)} jobs", "OK")
                            all_jobs.extend(ind_google_jobs)
                            tier_stats["Indeed-Google"] += len(ind_google_jobs)

                    delay(2.0, 4.0)

    # Print tier summary
    log("\n" + "-" * 50)
    log("TIER BREAKDOWN:")
    for tier, count in sorted(tier_stats.items()):
        if count > 0:
            log(f"   {tier:20s}: {count} jobs")
    log(f"   {'TOTAL':20s}: {len(all_jobs)} jobs (before dedup)")
    log("-" * 50)

    return all_jobs


def sync_to_google_sheets(df: pd.DataFrame, chunk_size: int = 100, max_retries: int = 3):
    """Sync job dataframe to Google Sheets Webhook with per-batch retries and resilience."""
    webhook_url = os.environ.get("GOOGLE_SHEETS_WEBHOOK_URL", "").strip()
    if not webhook_url and CONFIG_PATH.exists():
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                cfg = json.load(f)
                if not cfg.get("enabled", True) or not cfg.get("sync_on_crawl", True):
                    return
                webhook_url = cfg.get("webhook_url", "").strip()
        except Exception:
            pass

    if not webhook_url:
        return

    log("\nSyncing data to Google Sheets …", "INFO")
    clean_jobs = df.fillna("").to_dict(orient="records")
    total_jobs = len(clean_jobs)
    total_batches = (total_jobs + chunk_size - 1) // chunk_size
    success_batches = 0
    failed_batches = 0

    for i in range(0, total_jobs, chunk_size):
        batch_num = i // chunk_size + 1
        chunk = clean_jobs[i:i + chunk_size]
        payload = {
            "action": "sync_jobs",
            "jobs": chunk
        }

        batch_synced = False
        for attempt in range(1, max_retries + 1):
            try:
                log(f"   Syncing batch {batch_num}/{total_batches} ({len(chunk)} jobs){f' [attempt {attempt}/{max_retries}]' if attempt > 1 else ''} …", "INFO")
                res = requests.post(webhook_url, json=payload, timeout=60)
                if res.status_code == 200:
                    log(f"   Batch {batch_num}/{total_batches} synced → {res.text[:100]}", "OK")
                    batch_synced = True
                    success_batches += 1
                    break
                else:
                    log(f"   Batch {batch_num} failed HTTP {res.status_code}: {res.text[:100]}", "WARN")
            except requests.exceptions.RequestException as req_err:
                log(f"   Batch {batch_num} attempt {attempt} error: {req_err}", "WARN")

            if attempt < max_retries:
                time.sleep(attempt * 2)

        if not batch_synced:
            failed_batches += 1
            log(f"   Batch {batch_num} failed after {max_retries} attempts — continuing next batch.", "WARN")
        else:
            time.sleep(0.5)

    if failed_batches == 0:
        log(f"Google Sheets sync complete: all {total_jobs} jobs ({success_batches}/{total_batches} batches) synced successfully.", "OK")
    else:
        log(f"Google Sheets sync completed with {failed_batches} failed batches ({success_batches}/{total_batches} succeeded).", "WARN")


def main():
    parser = argparse.ArgumentParser(description="Indeed & Glassdoor Job Crawler -> Excel/CSV/Word (Browser-Primary)")
    parser.add_argument("--no-headless", action="store_true",
                        help="Show browser window (for CAPTCHA solving or visual verification)")
    parser.add_argument("--countries", nargs="*", default=None,
                        help="Country/City names to scrape (e.g. london toronto dubai amsterdam)")
    parser.add_argument("--sources", nargs="*", default=["indeed", "glassdoor"],
                        help="Sources to scrape: indeed, glassdoor (default: both)")
    parser.add_argument("--max-per-keyword", type=int, default=50,
                        help="Max jobs per keyword per source (default: 50)")
    parser.add_argument("--days", type=int, default=2,
                        help="Number of past days to crawl (default: 2)")
    parser.add_argument("--skip-google", action="store_true",
                        help="Skip Google Search fallback")
    parser.add_argument("--skip-browser", action="store_true",
                        help="Skip browser crawl (API-only mode)")
    args = parser.parse_args()

    headless = not args.no_headless
    sources = [s.lower() for s in args.sources]
    countries_filter = [c.lower() for c in args.countries] if args.countries else None

    all_jobs = run_crawler(headless, countries_filter, sources, args.max_per_keyword,
                           args.days, skip_google=args.skip_google, skip_browser=args.skip_browser)

    if not all_jobs:
        if MASTER_CSV_PATH.exists():
            log("No new live jobs extracted -- loading existing master database …", "WARN")
            jobs_df = pd.read_csv(MASTER_CSV_PATH, on_bad_lines="skip")
        else:
            log("No job records found to write.", "ERROR")
            return
    else:
        jobs_df = save_master_csv(all_jobs)

    if not jobs_df.empty:
        save_multi_sheet_excel(jobs_df)
        generate_word_report(jobs_df)
        sync_to_google_sheets(jobs_df)

    log("\n" + "=" * 65)
    log("  INDEED & GLASSDOOR CRAWL COMPLETE", "OK")
    log(f"   Total Unique Jobs : {len(jobs_df)}")
    log(f"   Master Excel      : {MASTER_EXCEL_PATH}")
    log(f"   Master CSV        : {MASTER_CSV_PATH}")
    log(f"   Word Report       : {DOCX_PATH}")
    log("=" * 65)


if __name__ == "__main__":
    main()
