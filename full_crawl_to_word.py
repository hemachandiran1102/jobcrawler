"""
full_crawl_to_word.py — 100% LinkedIn Job Crawler -> Word Document  v2
======================================================================
Uses the LinkedIn Voyager API (internal JSON API used by the LinkedIn
mobile/web app) with your saved browser session cookies.
No CSS selectors -> immune to layout changes.

How it works
------------
1. Launches Playwright to load your saved session (auto-login if needed)
2. Extracts the CSRF token + li_at cookie from the browser
3. Makes direct LinkedIn Voyager API calls (JSON) — much faster & reliable
4. Paginates through ALL results (25 per request) until exhausted
5. Covers 4 keywords × 10 countries
6. Writes: Word document + CSV + Excel

Usage
─────
  python full_crawl_to_word.py                        # headless, all countries
  python full_crawl_to_word.py --no-headless          # show browser (for login/CAPTCHA)
  python full_crawl_to_word.py --countries netherlands ireland
  python full_crawl_to_word.py --max-per-keyword 100  # cap per keyword
"""

import os
import sys
import time
import re
import json
import random
import hashlib
import argparse
import requests
import pandas as pd
from pathlib import Path
from datetime import datetime, timedelta

try:
    from playwright.sync_api import sync_playwright, TimeoutError as PWTimeout
except ImportError:
    print("ERROR: playwright not installed.  Run: pip install playwright && playwright install chromium")
    sys.exit(1)

try:
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
except ImportError:
    print("ERROR: python-docx not installed.  Run: pip install python-docx")
    sys.exit(1)

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False


# ══════════════════════════════════════════════════════════════════════
# PATHS
# ══════════════════════════════════════════════════════════════════════
_SCRIPT_DIR = Path(__file__).parent.resolve()
WORK_DIR    = Path(os.environ.get("WORK_DIR", str(_SCRIPT_DIR)))

# Session dir: prefer env var, then local JobCrawler .playwright-session, then job-automation
_LOCAL_SESSION = _SCRIPT_DIR / ".playwright-session"
_JOB_AUTO      = Path.home() / "Documents" / "Codex" / "2026-07-20" / "job-automation" / ".playwright-session"
SESSION_DIR    = Path(os.environ.get("SESSION_DIR", str(_LOCAL_SESSION if _LOCAL_SESSION.exists() else _JOB_AUTO)))

timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
RUN_DIR     = WORK_DIR / f"crawl_{timestamp}"
CSV_PATH    = RUN_DIR / f"full_crawl_jobs_{timestamp}.csv"
DOCX_PATH   = RUN_DIR / f"LinkedIn_FULL_Crawl_All_Countries_{timestamp}.docx"
EXCEL_PATH  = RUN_DIR / f"full_crawl_jobs_{timestamp}.xlsx"

# Single Master Spreadsheet Paths
MASTER_CSV_PATH   = WORK_DIR / "full_crawl_jobs.csv"
MASTER_EXCEL_PATH = WORK_DIR / "full_crawl_jobs.xlsx"
CONFIG_PATH       = WORK_DIR / "google_sheets_config.json"


# ══════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ══════════════════════════════════════════════════════════════════════
TARGET_COUNTRIES = [
    {"name": "Netherlands", "flag": "🇳🇱", "tier": 1, "geo_id": "102890719",
     "resume": "resumes/Hemachandiran_Giri_CV_NETHERLANDS.docx"},
    {"name": "Ireland",     "flag": "🇮🇪", "tier": 1, "geo_id": "104738515",
     "resume": "resumes/Hemachandiran_Giri_CV_IRELAND.docx"},
    {"name": "Sweden",      "flag": "🇸🇪", "tier": 1, "geo_id": "105117694",
     "resume": "resumes/Hemachandiran_Giri_CV_SWEDEN.docx"},
    {"name": "Denmark",     "flag": "🇩🇰", "tier": 1, "geo_id": "104514075",
     "resume": "resumes/Hemachandiran_Giri_CV_DENMARK.docx"},
    {"name": "Finland",     "flag": "🇫🇮", "tier": 1, "geo_id": "100456013",
     "resume": "resumes/Hemachandiran_Giri_CV_FINLAND.docx"},
    {"name": "France",      "flag": "🇫🇷", "tier": 2, "geo_id": "105015875",
     "resume": "resumes/Hemachandiran_Giri_CV_FRANCE.docx"},
    {"name": "Portugal",    "flag": "🇵🇹", "tier": 2, "geo_id": "100364837",
     "resume": "resumes/Hemachandiran_Giri_CV_PORTUGAL.docx"},
    {"name": "Poland",      "flag": "🇵🇱", "tier": 2, "geo_id": "105072130",
     "resume": "resumes/Hemachandiran_Giri_CV_POLAND.docx"},
    {"name": "Belgium",     "flag": "🇧🇪", "tier": 2, "geo_id": "100565514",
     "resume": "resumes/Hemachandiran_Giri_CV_BELGIUM.docx"},
    {"name": "Austria",     "flag": "🇦🇹", "tier": 2, "geo_id": "103883259",
     "resume": "resumes/Hemachandiran_Giri_CV_AUSTRIA.docx"},
    {"name": "Australia",   "flag": "🇦🇺", "tier": 1, "geo_id": "101452733",
     "resume": "resumes/Hemachandiran_Giri_CV_AUSTRALIA.docx"},
    {"name": "Singapore",   "flag": "🇸🇬", "tier": 1, "geo_id": "102454443",
     "resume": "resumes/Hemachandiran_Giri_CV_SINGAPORE.docx"},
    {"name": "Malaysia",    "flag": "🇲🇾", "tier": 2, "geo_id": "106808692",
     "resume": "resumes/Hemachandiran_Giri_CV_MALAYSIA.docx"},
    {"name": "New Zealand", "flag": "🇳🇿", "tier": 1, "geo_id": "105490917",
     "resume": "resumes/Hemachandiran_Giri_CV_NEW_ZEALAND.docx"},
]

SEARCH_TERMS = [
    "DevOps Engineer",
    "Cloud Engineer",
    "Site Reliability Engineer",
    "Platform Engineer",
]

VISA_KEYWORDS = [
    "visa", "sponsorship", "relocation", "work permit", "blue card",
    "kennismigrant", "critical skills", "passeport talent", "tech visa",
    "red-white-red", "tier 2", "skilled worker", "immigration",
]

SKILL_TAGS = [
    "AWS", "Kubernetes", "Docker", "Terraform", "CI/CD", "Python", "Linux",
    "GCP", "Azure", "Ansible", "Golang", "Kafka", "Jenkins", "ArgoCD",
    "Helm", "GitOps", "Prometheus", "Grafana", "EKS", "ECS", "Vault",
    "Pulumi", "Bash", "GitHub Actions", "GitLab CI", "Datadog",
]


# ══════════════════════════════════════════════════════════════════════
# UTILITIES
# ══════════════════════════════════════════════════════════════════════
def log(msg, level="INFO"):
    ts   = datetime.now().strftime("%H:%M:%S")
    icon = {"INFO": "[i]", "OK": "[OK]", "WARN": "[!]", "ERROR": "[ERR]"}.get(level, "")
    # Strip non-ASCII for Windows console compatibility
    safe = msg.encode("ascii", errors="replace").decode("ascii")
    print(f"[{ts}] {icon}  {safe}", flush=True)


def delay(lo=1.0, hi=2.0):
    time.sleep(random.uniform(lo, hi))


def is_authenticated_linkedin_url(url: str) -> bool:
    """Return True only for an authenticated LinkedIn page, never an authwall redirect."""
    value = (url or "").lower()
    return (
        "www.linkedin.com" in value
        and any(path in value for path in ("/feed", "/jobs", "/mynetwork"))
        and not any(marker in value for marker in ("/authwall", "/login", "/checkpoint", "/challenge"))
    )


def extract_visa(text: str) -> str:
    t = (text or "").lower()
    return "Yes" if any(k in t for k in VISA_KEYWORDS) else "No"


def extract_skills(text: str) -> str:
    t = (text or "").lower()
    found = [s for s in SKILL_TAGS if s.lower() in t]
    return ", ".join(found[:10]) if found else "N/A"


def normalise_posted_date(value) -> str:
    """Return a dashboard-friendly ISO date from LinkedIn/API date values.

    LinkedIn can return an ISO timestamp, a card label such as ``3 days ago``,
    or both (``2026-08-01 (3 days ago)``).  Keep a usable ISO date whenever
    the source provides one; use ``Unknown`` only when there is no date.
    """
    raw = str(value or "").strip()
    if not raw or raw.lower() in {"n/a", "unknown", "past week", "past month"}:
        return "Unknown"

    iso = re.search(r"\b(20\d{2}-\d{2}-\d{2})", raw)
    if iso:
        return iso.group(1)

    text = raw.lower().replace("reposted", "").strip()
    now = datetime.now()
    if text in {"today", "just now", "new"}:
        return now.strftime("%Y-%m-%d")
    if text == "yesterday":
        return (now - timedelta(days=1)).strftime("%Y-%m-%d")

    relative = re.search(r"(\d+)\s*(minute|hour|day|week|month)s?\s+ago", text)
    if relative:
        amount, unit = int(relative.group(1)), relative.group(2)
        days = {"minute": 0, "hour": 0, "day": amount, "week": amount * 7,
                "month": amount * 30}[unit]
        return (now - timedelta(days=days)).strftime("%Y-%m-%d")

    # Preserve non-standard source text rather than incorrectly inventing a date.
    return raw


def dedup_key(company: str, title: str) -> str:
    return hashlib.md5(
        f"{str(company).strip().lower()}|{str(title).strip().lower()}".encode()
    ).hexdigest()


# ══════════════════════════════════════════════════════════════════════
# PHASE 1 — BROWSER: get session cookies + CSRF token
# ══════════════════════════════════════════════════════════════════════
def get_linkedin_session(headless=True):
    """
    Get LinkedIn session tokens.
    Priority:
      1. .session.json file (created by save_session.py) — no browser needed
      2. Playwright persistent context (existing saved session)
    Returns (li_at, csrf_token, cookies_dict) or (None, None, None).
    """
    # ── Try .session.json first (fastest, no browser) ─────────────────
    session_json = WORK_DIR / ".session.json"
    if session_json.exists() and not os.environ.get("LINKEDIN_USE_BROWSER_SESSION"):
        try:
            data  = json.loads(session_json.read_text())
            li_at = data.get("li_at", "").strip()
            if li_at and len(li_at) > 50:
                log(f"Loaded li_at from {session_json}  (length {len(li_at)})", "OK")
                cookies = {"li_at": li_at}
                return li_at, "", cookies
        except Exception as e:
            log(f"Could not read .session.json: {e}", "WARN")

    # ── Try Playwright persistent context ─────────────────────────────
    log("Starting browser to extract session …")
    SESSION_DIR.mkdir(parents=True, exist_ok=True)

    with sync_playwright() as p:
        browser = p.chromium.launch_persistent_context(
            user_data_dir=str(SESSION_DIR),
            headless=headless,
            args=[
                "--no-sandbox",
                "--disable-dev-shm-usage",
                "--disable-blink-features=AutomationControlled",
            ],
            user_agent=(
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/125.0.0.0 Safari/537.36"
            ),
            viewport={"width": 1440, "height": 900},
            locale="en-US",
        )

        page = browser.new_page()

        # Navigate to LinkedIn
        log("Navigating to LinkedIn feed …")
        page.goto("https://www.linkedin.com/feed/", wait_until="domcontentloaded", timeout=45_000)
        delay(2, 4)

        # If not logged in, try credentials or wait for manual login
        if not is_authenticated_linkedin_url(page.url):
            email    = os.environ.get("LINKEDIN_EMAIL", "")
            password = os.environ.get("LINKEDIN_PASSWORD", "")

            if email and password:
                log("Logging in with credentials …")
                page.goto("https://www.linkedin.com/login", wait_until="domcontentloaded", timeout=30_000)
                delay(1, 2)
                page.fill("#username", email)
                delay(0.5, 1.0)
                page.fill("#password", password)
                delay(0.4, 0.8)
                page.click('button[type="submit"]')
                delay(4, 8)
            else:
                if not headless:
                    log("No credentials set. Please log in manually in the browser window.")
                    log("Waiting up to 3 minutes …", "WARN")
                    try:
                        page.wait_for_url("**/feed/**", timeout=180_000)
                    except PWTimeout:
                        log("Manual login timeout.", "ERROR")
                        browser.close()
                        return None, None, {}
                else:
                    log("No credentials and headless mode — cannot log in.", "ERROR")
                    browser.close()
                    return None, None, {}

        # Handle security challenge
        if any(x in page.url for x in ("checkpoint", "challenge", "captcha")):
            if not headless:
                log("Security challenge — please complete it in the browser.", "WARN")
                try:
                    page.wait_for_url("**/feed/**", timeout=180_000)
                except PWTimeout:
                    log("Challenge timeout.", "ERROR")
                    browser.close()
                    return None, None, {}
            else:
                log("Security challenge in headless mode — run with --no-headless.", "ERROR")
                browser.close()
                return None, None, {}

        if not is_authenticated_linkedin_url(page.url):
            log(f"Login check failed. URL: {page.url}", "ERROR")
            browser.close()
            return None, None, {}

        log("Logged in! Extracting session tokens …", "OK")

        # Extract cookies
        all_cookies = browser.cookies()
        cookies_dict = {c["name"]: c["value"] for c in all_cookies
                        if "linkedin.com" in c.get("domain", "")}

        li_at = cookies_dict.get("li_at", "")
        # CSRF token is in the JSESSIONID cookie or the lidc cookie
        csrf  = cookies_dict.get("JSESSIONID", "").strip('"')

        if not li_at:
            log("Could not extract li_at cookie — session may be invalid.", "ERROR")
            browser.close()
            return None, None, {}

        if not csrf:
            # Try to get it from the page meta tag
            try:
                csrf = page.evaluate(
                    "() => document.querySelector('meta[name=\"pageKey\"]')?.content || ''"
                )
            except Exception:
                csrf = ""

        log(f"li_at extracted (length {len(li_at)})", "OK")
        log(f"CSRF token: {csrf[:20]}…" if csrf else "CSRF: (empty — will try without)")

        browser.close()
        return li_at, csrf, cookies_dict


# ══════════════════════════════════════════════════════════════════════
# PHASE 2 — VOYAGER API: search jobs
# ══════════════════════════════════════════════════════════════════════
VOYAGER_BASE = "https://www.linkedin.com/voyager/api"

def _make_session(li_at: str, csrf: str, cookies_dict: dict) -> requests.Session:
    s = requests.Session()
    if not csrf:
        csrf = "ajax:1234567890123456789"
    headers = {
        "User-Agent"        : "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
        "Accept"            : "application/json, text/plain, */*",
        "Accept-Language"   : "en-US,en;q=0.9",
        "x-restli-protocol-version": "2.0.0",
        "csrf-token"        : csrf,
        "Referer"           : "https://www.linkedin.com/jobs/search/",
    }
    s.headers.update(headers)
    
    # Set cookies
    if li_at:
        s.cookies.set("li_at", li_at, domain=".linkedin.com")
    s.cookies.set("JSESSIONID", f'"{csrf}"', domain=".linkedin.com")
    
    for name, value in cookies_dict.items():
        if name not in ("li_at", "JSESSIONID"):
            s.cookies.set(name, value, domain=".linkedin.com")
            
    return s


def search_jobs_voyager(session: requests.Session, keyword: str, geo_id: str,
                        start: int = 0, count: int = 25, time_posted_range: str = "r86400") -> dict:
    """
    Call LinkedIn Guest/API endpoint for the last 24 hours (r86400).
    """
    params = {
        "keywords": keyword,
        "location": "Europe",
        "geoId": geo_id,
        "f_TPR": time_posted_range,  # Past 24 hours
        "f_JT": "F",          # Full-time
        "sortBy": "DD",
        "start": start,
        "count": count
    }
    url = "https://www.linkedin.com/jobs-guest/jobs/api/seeMoreJobPostings/search"

    for attempt in range(3):
        try:
            resp = session.get(url, params=params, timeout=20, allow_redirects=False)
            if resp.status_code == 200:
                html = resp.text
                job_ids = list(dict.fromkeys(
                    re.findall(r'data-entity-urn="urn:li:jobPosting:(\d+)"', html) or
                    re.findall(r'data-job-id="(\d+)"', html) or
                    re.findall(r'/jobs/view/(\d+)', html) or
                    re.findall(r'urn:li:jobPosting:(\d+)', html)
                ))
                
                elements = []
                for jid in job_ids:
                    title = keyword
                    company = "N/A"
                    location = "Europe"
                    is_easy_apply = "No"
                    is_remote = "No"
                    
                    # Check for Easy Apply / Remote in HTML card block
                    m_block = re.search(rf'jobPosting:{jid}.*?</li>', html, re.DOTALL)
                    if m_block:
                        block_text = m_block.group(0)
                        m_t = re.search(r'<h3[^>]*>\s*(.*?)\s*</h3>', block_text, re.DOTALL)
                        if m_t:
                            title = re.sub(r'<[^>]+>', '', m_t.group(1)).strip()
                        m_c = re.search(r'<h4[^>]*>\s*(.*?)\s*</h4>', block_text, re.DOTALL)
                        if m_c:
                            company = re.sub(r'<[^>]+>', '', m_c.group(1)).strip()
                        m_l = re.search(r'job-search-card__location[^>]*>\s*(.*?)\s*</', block_text, re.DOTALL)
                        if m_l:
                            location = re.sub(r'<[^>]+>', '', m_l.group(1)).strip()
                        posted_date = "N/A"
                        m_time = re.search(r'<time[^>]*datetime=["\']([^"\']+)["\'][^>]*>(.*?)</time>', block_text, re.DOTALL)
                        if m_time:
                            dt_val = m_time.group(1).strip()
                            rel_txt = re.sub(r'<[^>]+>', '', m_time.group(2)).strip()
                            if dt_val and rel_txt:
                                posted_date = f"{dt_val} ({rel_txt})"
                            elif dt_val:
                                posted_date = dt_val
                            elif rel_txt:
                                posted_date = rel_txt

                        if "easy apply" in block_text.lower() or "simple apply" in block_text.lower():
                            is_easy_apply = "Yes"
                            if "remote" in block_text.lower() or "work from home" in block_text.lower():
                                is_remote = "Yes"
                                
                        elements.append({
                            "jobPostingUrn": f"urn:li:jobPosting:{jid}",
                            "title": title,
                            "location": location,
                            "companyName": company,
                            "job_id": jid,
                            "posted_date": posted_date,
                            "job_url": f"https://www.linkedin.com/jobs/view/{jid}",
                            "easy_apply": is_easy_apply,
                            "remote": is_remote
                        })
                return {"elements": elements}
            elif resp.status_code == 429:
                wait_sec = 2
                log(f"    Rate limited (429) — sleeping {wait_sec}s (attempt {attempt+1}/3) …", "WARN")
                time.sleep(wait_sec)
            else:
                log(f"    API returned {resp.status_code} for {keyword}", "WARN")
                return {}
        except Exception as e:
            log(f"    Request error: {e}", "WARN")
            time.sleep(2)
    return {}


def search_jobs_voyager_v2(session: requests.Session, keyword: str, geo_id: str,
                            start: int = 0, count: int = 25, time_posted_range: str = "r86400") -> dict:
    """
    Alternative Voyager endpoint used by LinkedIn web search page.
    """
    params = {
        "count"        : count,
        "filters"      : f"List(geoUrn:urn%3Ali%3Ageo%3A{geo_id},timePostedRange:{time_posted_range},workType:F)",
        "keywords"     : keyword,
        "origin"       : "JOB_SEARCH_PAGE_SEARCH_BUTTON",
        "q"            : "jserpFilters",
        "start"        : start,
    }
    url = f"{VOYAGER_BASE}/jobs/search/hits"
    try:
        resp = session.get(url, params=params, timeout=20, allow_redirects=False)
        if resp.status_code == 200:
            return resp.json()
        elif resp.status_code == 429:
            log(f"    Rate limited (429) v2 — sleeping 2s …", "WARN")
            time.sleep(2)
            return {}
        else:
            return {}
    except Exception as e:
        log(f"    Request error (v2): {e}", "WARN")
        return {}


def get_job_detail(session: requests.Session, job_id: str) -> dict:
    """Fetch full job description for a single job posting."""
    url = f"{VOYAGER_BASE}/jobs/jobPostings/{job_id}"
    try:
        resp = session.get(url, timeout=15)
        if resp.status_code == 200:
            return resp.json()
    except Exception:
        pass
    return {}


def parse_voyager_jobs(data: dict) -> list:
    """
    Parse job hits from Voyager API response.
    Returns list of dicts with basic job info.
    """
    jobs = []

    # The response structure varies — try multiple paths
    elements = (
        data.get("elements", [])
        or data.get("data", {}).get("elements", [])
        or []
    )

    # Also check included items (normalized format)
    included = data.get("included", [])
    included_jobs = {}
    for item in included:
        t = item.get("$type", "")
        if "JobPosting" in t:
            eid = item.get("entityUrn", "").split(":")[-1]
            if eid:
                included_jobs[eid] = item

    for elem in elements:
        try:
            # Get job ID from entity URN
            urn = (
                elem.get("trackingUrn", "")
                or elem.get("jobPostingUrn", "")
                or elem.get("entityUrn", "")
                or ""
            )
            job_id = urn.split(":")[-1] if ":" in urn else ""

            # Try to get job data from element or included map
            job_data = elem.get("jobPosting", {}) or included_jobs.get(job_id, {}) or elem

            title   = (
                job_data.get("title", "")
                or elem.get("title", "")
                or "N/A"
            )
            company = (
                job_data.get("companyDetails", {})
                        .get("company", {})
                        .get("name", "")
                or job_data.get("formattedEmploymentStatus", "")
                or elem.get("companyName", "")
                or "N/A"
            )
            location = (
                job_data.get("formattedLocation", "")
                or elem.get("location", "")
                or ""
            )
            # Guest-search cards already supply a posted_date.  Previously this
            # parser ignored it, so every guest-search result was saved as
            # "Unknown" despite the date being present in the source card.
            source_posted = (
                elem.get("posted_date")
                or job_data.get("posted_date")
                or job_data.get("postedDate")
                or elem.get("postedDate")
            )
            listed_at = (
                job_data.get("listedAt", 0)
                or job_data.get("originalListedAt", 0)
                or elem.get("listedAt", 0)
                or elem.get("originalListedAt", 0)
            )
            if source_posted:
                posted = normalise_posted_date(source_posted)
            elif listed_at:
                try:
                    posted = datetime.fromtimestamp(listed_at / 1000).strftime("%Y-%m-%d %H:%M")
                except Exception:
                    posted = "Unknown"
            else:
                posted = "Unknown"

            job_url = f"https://www.linkedin.com/jobs/view/{job_id}" if job_id else ""

            description = job_data.get("description", {})
            description_text = job_data.get("descriptionText", {})
            desc = (
                description.get("text", "") if isinstance(description, dict) else str(description)
            ) or (
                description_text.get("text", "") if isinstance(description_text, dict) else str(description_text)
            )

            if title and title != "N/A" and job_id:
                jobs.append({
                    "job_id"   : job_id,
                    "title"    : title,
                    "company"  : company,
                    "location" : location,
                    "posted"   : posted,
                    "job_url"  : job_url,
                    "desc"     : desc,
                    "easy_apply": elem.get("easy_apply") or job_data.get("easy_apply") or "No",
                    "remote"   : elem.get("remote") or job_data.get("remote") or "On-site / Hybrid",
                })
        except Exception:
            continue

    return jobs


# ══════════════════════════════════════════════════════════════════════
# PHASE 3 — Browser fallback scraper (improved selectors)
# ══════════════════════════════════════════════════════════════════════

def scrape_with_browser(headless: bool, countries_filter, max_per_kw: int) -> list:
    """
    Fallback: Use Playwright to scrape the Jobs search page directly.
    Tries multiple CSS selector strategies and waits for dynamic content.
    """
    log("Using browser-based scraping fallback …", "WARN")
    SESSION_DIR.mkdir(parents=True, exist_ok=True)
    all_jobs = []

    countries = TARGET_COUNTRIES
    if countries_filter:
        countries = [c for c in TARGET_COUNTRIES if c["name"].lower() in countries_filter]

    with sync_playwright() as p:
        browser = p.chromium.launch_persistent_context(
            user_data_dir=str(SESSION_DIR),
            headless=headless,
            args=["--no-sandbox", "--disable-dev-shm-usage",
                  "--disable-blink-features=AutomationControlled"],
            user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/125.0.0.0 Safari/537.36",
            viewport={"width": 1440, "height": 900},
            locale="en-US",
        )
        page = browser.new_page()

        # Login check
        page.goto("https://www.linkedin.com/feed/", wait_until="domcontentloaded", timeout=40_000)
        delay(2, 4)
        if not is_authenticated_linkedin_url(page.url):
            log("Not logged in — cannot do browser scraping.", "ERROR")
            browser.close()
            return []

        log("Logged in via browser.", "OK")

        for country in countries:
            for term in SEARCH_TERMS:
                kw_jobs  = []
                seen_ids = set()
                start    = 0

                while len(kw_jobs) < max_per_kw:
                    kw_enc  = term.replace(" ", "%20").replace("&", "%26")
                    loc_enc = country["name"].replace(" ", "%20")
                    geo     = country["geo_id"]
                    url = (
                        f"https://www.linkedin.com/jobs/search/"
                        f"?keywords={kw_enc}&location={loc_enc}"
                        f"&geoId={geo}&f_TPR=r1209600&f_JT=F&sortBy=DD&start={start}"
                    )
                    log(f"  [{country['name']}] '{term}' start={start}")

                    try:
                        page.goto(url, wait_until="domcontentloaded", timeout=35_000)
                        delay(3, 5)
                    except Exception as e:
                        log(f"  Navigation error: {e}", "WARN")
                        break

                    if "authwall" in page.url or "login" in page.url:
                        log("  Redirected to login — stopping.", "WARN")
                        break

                    # Scroll to trigger lazy loading
                    for _ in range(5):
                        page.keyboard.press("End")
                        delay(0.8, 1.5)
                    delay(1, 2)

                    # Try to get jobs from the page JSON (LinkedIn embeds state in <code> tags)
                    page_jobs = []
                    try:
                        code_els = page.query_selector_all("code")
                        for cel in code_els:
                            raw = cel.inner_text()
                            if '"jobPostings"' in raw or '"elements"' in raw:
                                try:
                                    jdata = json.loads(raw)
                                    parsed = parse_voyager_jobs(jdata)
                                    if parsed:
                                        page_jobs.extend(parsed)
                                except Exception:
                                    pass
                    except Exception:
                        pass

                    if not page_jobs:
                        # Fall back to DOM scraping
                        page_jobs = _dom_scrape(page, country["name"])

                    if not page_jobs:
                        log(f"  No jobs found at start={start} — end of results.")
                        break

                    new_count = 0
                    for j in page_jobs:
                        jid = j.get("job_id") or j.get("job_url", "")
                        if jid in seen_ids:
                            continue
                        seen_ids.add(jid)
                        kw_jobs.append(j)
                        new_count += 1
                        if len(kw_jobs) >= max_per_kw:
                            break

                    log(f"  +{new_count} new (total: {len(kw_jobs)})")
                    if new_count == 0:
                        break
                    start += 25
                    delay(2, 4)

                # Convert to standard format
                for j in kw_jobs:
                    all_jobs.append(_normalise_job(j, country, term))

                log(f"  ✅ {country['name']} / '{term}': {len(kw_jobs)} jobs", "OK")
                delay(3, 6)

        browser.close()

    return all_jobs


def _dom_scrape(page, country_name: str) -> list:
    """
    Multi-strategy scraper:
    1. Extract jobs from embedded JSON (LinkedIn inlines Voyager data in <code> tags)
    2. Fall back to CSS card selectors
    3. Fall back to regex on raw page source
    """
    jobs = []

    # ── Strategy 1: Embedded JSON in <code> tags ─────────────────────────
    try:
        code_els = page.query_selector_all("code")
        for cel in code_els:
            raw = cel.inner_text()
            if not raw or len(raw) < 50:
                continue
            if '"jobPosting"' in raw or '"jobTitle"' in raw or '"entityUrn"' in raw:
                try:
                    data = json.loads(raw)
                    parsed = parse_voyager_jobs(data)
                    for j in parsed:
                        if j.get("job_url") or j.get("job_id"):
                            jobs.append(j)
                except Exception:
                    pass
        if jobs:
            log(f"    Extracted {len(jobs)} jobs from embedded JSON")
            return jobs
    except Exception:
        pass

    # ── Strategy 2: Page source regex for job IDs ─────────────────────────
    try:
        src = page.content()
        # Find all /jobs/view/<id> patterns
        job_ids = list(dict.fromkeys(re.findall(r'/jobs/view/(\d+)', src)))
        if job_ids:
            log(f"    Found {len(job_ids)} job IDs in page source via regex")
            for jid in job_ids[:50]:  # cap at 50 from one page
                # Try to extract title + company near this ID in the source
                jobs.append({
                    "job_id"  : jid,
                    "title"   : "See LinkedIn",
                    "company" : "See LinkedIn",
                    "location": country_name,
                    "posted"  : "Unknown",
                    "job_url" : f"https://www.linkedin.com/jobs/view/{jid}",
                    "desc"    : "",
                })
            return jobs
    except Exception:
        pass

    # ── Strategy 3: CSS card selectors ────────────────────────────────────
    card_selectors = [
        "li.jobs-search-results__list-item",
        "div.job-card-container",
        "div.jobs-search-results-grid__item",
        "[data-job-id]",
        "li.occludable-update",
        ".job-card-list",
        "li[class*='jobs-search']",
        "div[class*='job-card']",
    ]
    cards = []
    for sel in card_selectors:
        try:
            cards = page.query_selector_all(sel)
            if cards:
                log(f"    CSS selector matched: {sel} ({len(cards)} cards)")
                break
        except Exception:
            continue

    for card in cards:
        try:
            job_id = (
                card.get_attribute("data-job-id") or ""
            )
            if not job_id:
                urn = card.get_attribute("data-entity-urn") or ""
                job_id = urn.split(":")[-1] if ":" in urn else ""

            title_el = card.query_selector(
                "a.job-card-list__title, "
                ".job-card-container__link, "
                "a[href*='/jobs/view/'], "
                ".job-card-list__title--link, "
                "strong, "
                "h3"
            )
            company_el = card.query_selector(
                ".job-card-container__primary-description, "
                ".job-card-container__company-name, "
                ".artdeco-entity-lockup__subtitle, "
                "span.t-14"
            )
            location_el = card.query_selector(
                ".job-card-container__metadata-item, "
                "li.job-card-container__metadata-wrapper, "
                ".artdeco-entity-lockup__caption, "
                "span.job-card-container__metadata-item"
            )

            title    = title_el.inner_text().strip()    if title_el    else "N/A"
            company  = company_el.inner_text().strip()  if company_el  else "N/A"
            location = location_el.inner_text().strip() if location_el else country_name

            href = ""
            if title_el:
                href = title_el.get_attribute("href") or ""
            if not href and job_id:
                href = f"https://www.linkedin.com/jobs/view/{job_id}"
            if href.startswith("/"):
                href = "https://www.linkedin.com" + href
            job_url = href.split("?")[0]

            if not job_url and not job_id:
                continue
            if not job_id:
                m = re.search(r"/jobs/view/(\d+)", job_url)
                job_id = m.group(1) if m else ""

            jobs.append({
                "job_id"  : job_id,
                "title"   : title,
                "company" : company,
                "location": location,
                "posted"  : "Past month",
                "job_url" : job_url or f"https://www.linkedin.com/jobs/view/{job_id}",
                "desc"    : "",
            })
        except Exception:
            continue

    return jobs


def _normalise_job(j: dict, country: dict, term: str) -> dict:
    desc = j.get("desc", "")
    title = j.get("title", "N/A")
    title_lower = title.lower()
    
    # Calculate match score based on candidate profile keywords
    score = 70
    if any(k in title_lower for k in ["senior", "lead", "principal", "architect"]):
        score += 15
    if any(k in title_lower for k in ["devops", "cloud", "sre"]):
        score += 15
        
    return {
        "Date"                      : datetime.now().strftime("%Y-%m-%d"),
        "Country"                   : country["name"],
        "Flag"                      : country["flag"],
        "Tier"                      : country["tier"],
        "Location"                  : j.get("location", country["name"]),
        "Company"                   : j.get("company", "N/A"),
        "Job Title"                 : title,
        "Search Keyword"            : term,
        "Job URL"                   : j.get("job_url", ""),
        "Posted Date"               : j.get("posted_date") or j.get("posted") or "Past week",
        "Easy Apply"                : j.get("easy_apply", "No"),
        "Remote / Workplace"        : j.get("remote", "On-site / Hybrid"),
        "Match Score"               : f"{score}%",
        "Visa Sponsorship Mentioned": extract_visa(desc),
        "Required Skills"           : extract_skills(desc),
        "Resume File Path"          : country["resume"],
        "Applied Status"            : "No",
        "Notes"                     : "",
    }


# ══════════════════════════════════════════════════════════════════════
# PHASE 4 — Orchestrate: Voyager API first, browser fallback
# ══════════════════════════════════════════════════════════════════════

def run_full_crawl(headless=True, countries_filter=None, max_per_kw=999, time_posted_range="r86400") -> list:
    log("=" * 65)
    log("🚀  FULL CRAWL — starting")
    log(f"   Session dir   : {SESSION_DIR}")
    log(f"   Output dir    : {WORK_DIR}")
    log(f"   Time window   : {time_posted_range or 'All time'}")
    log("=" * 65)

    countries = TARGET_COUNTRIES
    if countries_filter:
        countries = [c for c in TARGET_COUNTRIES if c["name"].lower() in countries_filter]

    # ── Step 1: get session tokens from browser ────────────────────
    li_at, csrf, cookies_dict = get_linkedin_session(headless=headless)

    if not li_at:
        log("Could not get LinkedIn session. Trying browser scraping directly …", "WARN")
        return scrape_with_browser(headless, countries_filter, max_per_kw)

    # ── Step 2: call Voyager API ───────────────────────────────────
    session  = _make_session(li_at, csrf, cookies_dict)
    all_jobs = []

    for country in countries:
        log(f"\n{'─'*60}")
        log(f"{country['flag']}  Crawling {country['name']} via Voyager API …")
        country_seen = set()

        for term in SEARCH_TERMS:
            log(f"  🔍  Keyword: \"{term}\"")
            kw_jobs = []
            start   = 0

            while len(kw_jobs) < max_per_kw:
                log(f"    start={start} …")
                data = search_jobs_voyager(session, term, country["geo_id"], start, time_posted_range=time_posted_range)

                # If main endpoint gives nothing, try alternative
                if not data or not data.get("elements"):
                    data = search_jobs_voyager_v2(session, term, country["geo_id"], start, time_posted_range=time_posted_range)

                if not data:
                    break

                parsed = parse_voyager_jobs(data)
                if not parsed:
                    log(f"    No jobs in response at start={start}")
                    break

                new_count = 0
                for j in parsed:
                    jid = j.get("job_url", "") or j.get("job_id", "")
                    dk  = dedup_key(j.get("company",""), j.get("title",""))
                    if jid in country_seen or dk in country_seen:
                        continue
                    country_seen.add(jid)
                    country_seen.add(dk)
                    kw_jobs.append(j)
                    new_count += 1
                    if len(kw_jobs) >= max_per_kw:
                        break

                log(f"    +{new_count} jobs (total this keyword: {len(kw_jobs)})")
                if new_count == 0:
                    break

                start += 25
                delay(1.0, 2.0)

            log(f"  ✅  \"{term}\" → {len(kw_jobs)} jobs", "OK")

            for j in kw_jobs:
                all_jobs.append(_normalise_job(j, country, term))

            delay(1, 2)

        log(f"{country['flag']}  {country['name']} complete.", "OK")

    # ── If Voyager gave nothing, fall back to browser scraping ─────
    if not all_jobs:
        log("Voyager API returned no results — falling back to browser scraping.", "WARN")
        all_jobs = scrape_with_browser(headless, countries_filter, max_per_kw)

    log(f"\n📊 LinkedIn total: {len(all_jobs)}")
    return all_jobs


# ══════════════════════════════════════════════════════════════════════
# PHASE 5 — Free API supplements
# ══════════════════════════════════════════════════════════════════════

def fetch_free_apis() -> list:
    log("\nFetching free API supplements …")
    api_jobs   = []
    name_lower = {c["name"].lower() for c in TARGET_COUNTRIES}
    resume_map = {c["name"].lower(): c["resume"] for c in TARGET_COUNTRIES}
    flag_map   = {c["name"].lower(): c["flag"]   for c in TARGET_COUNTRIES}
    tier_map   = {c["name"].lower(): c["tier"]   for c in TARGET_COUNTRIES}
    kw_filter  = {"devops", "cloud", "sre", "platform", "infrastructure",
                  "kubernetes", "devsecops", "reliability", "mlops"}

    def _kw_match(title: str) -> bool:
        return any(k in title.lower() for k in kw_filter)

    def _country_match(location: str):
        loc = location.lower()
        for cn in name_lower:
            if cn in loc:
                return cn
        return None

    # Arbeitnow
    try:
        r = requests.get("https://www.arbeitnow.com/api/job-board-api", timeout=20)
        if r.ok:
            before = len(api_jobs)
            for job in r.json().get("data", []):
                cn = _country_match(job.get("location", ""))
                if not cn or not _kw_match(job.get("title", "")):
                    continue
                desc = job.get("description", "")
                api_jobs.append({
                    "Date": datetime.now().strftime("%Y-%m-%d"),
                    "Country": cn.title(), "Flag": flag_map.get(cn, ""),
                    "Tier": tier_map.get(cn, 2),
                    "Location": job.get("location", ""),
                    "Company": job.get("company_name", "N/A"),
                    "Job Title": job.get("title", ""),
                    "Search Keyword": "Free API",
                    "Job URL": job.get("url", ""),
                    "Posted Date": str(job.get("created_at", ""))[:10],
                    "Visa Sponsorship Mentioned": "Yes" if job.get("visa_sponsorship") else extract_visa(desc),
                    "Required Skills": ", ".join(job.get("tags", [])[:8]) or extract_skills(desc),
                    "Resume File Path": resume_map.get(cn, ""),
                    "Applied Status": "No", "Notes": "Source: Arbeitnow",
                })
            log(f"  Arbeitnow: {len(api_jobs) - before} jobs")
    except Exception as e:
        log(f"  Arbeitnow error: {e}", "WARN")

    # Remotive
    before = len(api_jobs)
    try:
        r = requests.get("https://remotive.com/api/remote-jobs?category=devops", timeout=20)
        if r.ok:
            for job in r.json().get("jobs", []):
                if not _kw_match(job.get("title", "")):
                    continue
                loc = job.get("candidate_required_location", "")
                cn  = _country_match(loc)
                if not cn:
                    if any(x in loc.lower() for x in ("europe", "eu", "worldwide", "emea")):
                        cn = "netherlands"
                    else:
                        continue
                desc = job.get("description", "")
                api_jobs.append({
                    "Date": datetime.now().strftime("%Y-%m-%d"),
                    "Country": cn.title(), "Flag": flag_map.get(cn, ""),
                    "Tier": tier_map.get(cn, 2), "Location": loc or "Remote EU",
                    "Company": job.get("company_name", "N/A"),
                    "Job Title": job.get("title", ""),
                    "Search Keyword": "Free API",
                    "Job URL": job.get("url", ""),
                    "Posted Date": str(job.get("publication_date", ""))[:10],
                    "Visa Sponsorship Mentioned": extract_visa(desc),
                    "Required Skills": ", ".join(job.get("tags", [])[:8]) or extract_skills(desc),
                    "Resume File Path": resume_map.get(cn, ""),
                    "Applied Status": "No", "Notes": "Source: Remotive",
                })
        log(f"  Remotive: {len(api_jobs) - before} jobs")
    except Exception as e:
        log(f"  Remotive error: {e}", "WARN")

    log(f"  Total API jobs: {len(api_jobs)}", "OK")
    return api_jobs


# ══════════════════════════════════════════════════════════════════════
# PHASE 6 — Save CSV
# ══════════════════════════════════════════════════════════════════════

def parse_sort_datetime(posted_val):
    s = str(posted_val or "").strip()
    if not s or s in ["N/A", "Unknown", "Past week", "Past 3 weeks", "Past 2 weeks", "Past month"]:
        return datetime(2000, 1, 1)
    try:
        m = re.search(r'(\d{4}-\d{2}-\d{2})(?:\s+(\d{2}:\d{2}))?', s)
        if m:
            d = m.group(1)
            t = m.group(2) or "00:00"
            return datetime.strptime(f"{d} {t}", "%Y-%m-%d %H:%M")
        now = datetime.now()
        m_rel = re.search(r'(\d+)\s*(min|minute|hour|hr|day|d|week|wk)', s.lower())
        if m_rel:
            num = int(m_rel.group(1))
            unit = m_rel.group(2)
            if 'min' in unit:
                return now - timedelta(minutes=num)
            elif 'hour' in unit or 'hr' in unit:
                return now - timedelta(hours=num)
            elif 'day' in unit or 'd' in unit:
                return now - timedelta(days=num)
            elif 'week' in unit or 'wk' in unit:
                return now - timedelta(weeks=num)
    except Exception:
        pass
    return datetime(2000, 1, 1)


def extract_city(loc_str):
    if not loc_str or loc_str in ["N/A", "Europe", "Unknown"]:
        return "Other Locations"
    parts = [p.strip() for p in str(loc_str).split(",")]
    if len(parts) >= 1:
        city = parts[0]
        city = re.sub(r'^(Greater|Metropolitan|City of)\s+', '', city, flags=re.I).strip()
        return city if city else "Other Locations"
    return "Other Locations"


def sync_to_google_sheets(df: pd.DataFrame):
    """Sync job dataframe to Google Sheets Webhook if configured."""
    webhook_url = os.environ.get("GOOGLE_SHEETS_WEBHOOK_URL", "").strip()
    if not webhook_url and CONFIG_PATH.exists():
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                cfg = json.load(f)
                webhook_url = cfg.get("webhook_url", "").strip()
        except Exception:
            pass

    if not webhook_url:
        log("Google Sheets webhook URL not configured (skipping sync).", "INFO")
        return

    log("\nSyncing data to Google Sheets …", "INFO")
    try:
        clean_jobs = df.fillna("").to_dict(orient="records")
        total_jobs = len(clean_jobs)
        chunk_size = 250
        
        for i in range(0, total_jobs, chunk_size):
            chunk = clean_jobs[i:i + chunk_size]
            payload = {
                "action": "sync_jobs",
                "jobs": chunk
            }
            log(f"   Syncing batch {i // chunk_size + 1}/{(total_jobs + chunk_size - 1) // chunk_size} ({len(chunk)} jobs) …", "INFO")
            res = requests.post(webhook_url, json=payload, timeout=30)
            if res.status_code == 200:
                log(f"   Batch synced → {res.text[:100]}", "OK")
            else:
                log(f"   Batch failed HTTP {res.status_code}: {res.text[:100]}", "WARN")
            time.sleep(0.5)
            
        log(f"Google Sheets sync complete: {total_jobs} total jobs processed.", "OK")
    except Exception as err:
        log(f"Google Sheets sync error: {err}", "WARN")


def save_csv(all_jobs: list) -> pd.DataFrame:
    today_str = datetime.now().strftime("%Y-%m-%d")
    
    # Tag incoming jobs with Crawl Date and dedup keys
    incoming_jobs = []
    for job in all_jobs:
        j = dict(job)
        if not j.get("Crawl Date") and not j.get("Date"):
            j["Crawl Date"] = today_str
            j["Date"] = today_str
        elif j.get("Crawl Date") and not j.get("Date"):
            j["Date"] = j["Crawl Date"]
        elif j.get("Date") and not j.get("Crawl Date"):
            j["Crawl Date"] = j["Date"]
        incoming_jobs.append(j)

    new_df = pd.DataFrame(incoming_jobs)
    
    # Merge with Master CSV if it exists
    if MASTER_CSV_PATH.exists():
        try:
            log(f"Loading existing master sheet → {MASTER_CSV_PATH}", "INFO")
            master_df = pd.read_csv(MASTER_CSV_PATH, on_bad_lines="skip")
            if not master_df.empty:
                # Build map of existing master records by URL / Dedup key
                master_records = master_df.to_dict(orient="records")
                existing_map = {}
                for m in master_records:
                    url = str(m.get("Job URL", "")).strip()
                    comp = str(m.get("Company", "")).strip().lower()
                    title = str(m.get("Job Title", "")).strip().lower()
                    key = url if url else f"{comp}|{title}"
                    existing_map[key] = m
                
                # Update / Append incoming jobs into existing master map
                for item in incoming_jobs:
                    url = str(item.get("Job URL", "")).strip()
                    comp = str(item.get("Company", "")).strip().lower()
                    title = str(item.get("Job Title", "")).strip().lower()
                    key = url if url else f"{comp}|{title}"
                    
                    if key in existing_map:
                        prev = existing_map[key]
                        # Retain existing user fields if already marked
                        if str(prev.get("Applied Status", "")).strip().lower() == "yes":
                            item["Applied Status"] = "Yes"
                        if prev.get("Notes"):
                            item["Notes"] = prev.get("Notes")
                        if prev.get("Crawl Date"):
                            item["Crawl Date"] = prev.get("Crawl Date")
                            item["Date"] = prev.get("Crawl Date")
                    existing_map[key] = item

                merged_jobs = list(existing_map.values())
                df = pd.DataFrame(merged_jobs)
            else:
                df = new_df
        except Exception as e:
            log(f"Error reading master CSV: {e} — using incoming dataset", "WARN")
            df = new_df
    else:
        df = new_df

    if df.empty:
        return df

    df.drop_duplicates(subset=["Job URL"], keep="last", inplace=True)
    df["_dk"] = df.apply(lambda r: dedup_key(r.get("Company",""), r.get("Job Title","")), axis=1)
    df.drop_duplicates(subset=["_dk"], keep="first", inplace=True)
    df.drop(columns=["_dk"], inplace=True)
    
    # Ensure Crawl Date exists
    if "Crawl Date" not in df.columns:
        df["Crawl Date"] = today_str
    if "Date" not in df.columns:
        df["Date"] = df["Crawl Date"]

    # Extract City and calculate City Openings popularity count
    df["City"] = df["Location"].apply(extract_city)
    city_counts = df["City"].value_counts().to_dict()
    df["City Openings"] = df["City"].map(city_counts).fillna(0).astype(int)

    # Sort by City Openings ASCENDING and Posted Date DESCENDING
    if "Posted Date" in df.columns:
        df["_sort_dt"] = df["Posted Date"].apply(parse_sort_datetime)
        df.sort_values(by=["City Openings", "_sort_dt"], ascending=[True, False], inplace=True)
        df.drop(columns=["_sort_dt"], inplace=True)
    else:
        df.sort_values(by=["City Openings"], ascending=True, inplace=True)

    df.reset_index(drop=True, inplace=True)

    # Write to Single Master CSV at root and timestamped CSV in run dir
    df.to_csv(MASTER_CSV_PATH, index=False)
    df.to_csv(CSV_PATH, index=False)
    log(f"Master CSV updated → {MASTER_CSV_PATH} ({len(df)} total jobs maintaining single master sheet)", "OK")
    log(f"Run CSV archived  → {CSV_PATH}", "OK")

    # Sync to Google Sheets
    sync_to_google_sheets(df)

    return df


# ══════════════════════════════════════════════════════════════════════
# PHASE 7 — Generate Word Document
# ══════════════════════════════════════════════════════════════════════

def _shade(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:color="auto" w:val="clear"/>'
    ))


def _hdr_cell(cell, text: str, bg="1F4E79", width=None):
    if width:
        cell.width = width
    _shade(cell, bg)
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _hyperlink(para, url: str, text: str):
    part = para.part
    rid  = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    h = parse_xml(
        f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" r:id="{rid}"/>'
    )
    r = parse_xml('<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
    rpr = parse_xml(
        '<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:color w:val="1155CC"/><w:u w:val="single"/></w:rPr>'
    )
    t = parse_xml(
        f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' xml:space="preserve">{text}</w:t>'
    )
    r.append(rpr); r.append(t); h.append(r)
    para._p.append(h)


def _section(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(6)


def generate_word(jobs_df: pd.DataFrame):
    log("\nGenerating Word document …")
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Inches(0.7)
        sec.bottom_margin = Inches(0.7)
        sec.left_margin   = Inches(0.8)
        sec.right_margin  = Inches(0.8)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    run_date = datetime.now().strftime("%d %B %Y  %H:%M")
    total    = len(jobs_df)

    # ── Title ────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("LinkedIn Cloud & DevOps Jobs — Full Crawl Report")
    r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0, 51, 102)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run("All Countries · All Pages · Authenticated Session")
    r.font.size = Pt(12); r.font.color.rgb = RGBColor(31, 78, 121)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    r = p.add_run(
        f"Candidate: Hemachandiran Giri  ·  AWS Certified DevOps Professional\n"
        f"Total Jobs: {total}  ·  Generated: {run_date}"
    )
    r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(100, 100, 100)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)

    # ── Country Dashboard ────────────────────────────────────────────
    _section(doc, "1.  Country Dashboard")

    dash = doc.add_table(rows=1, cols=5)
    dash.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdrs = ["Country", "Flag", "Total Jobs", "Visa Mention", "Resume"]
    wids = [Inches(1.6), Inches(0.5), Inches(1.1), Inches(1.1), Inches(2.8)]
    for i, (h, w) in enumerate(zip(hdrs, wids)):
        _hdr_cell(dash.rows[0].cells[i], h, width=w)

    v_col = "Visa Sponsorship Mentioned"
    for country in TARGET_COUNTRIES:
        cn   = country["name"]
        rows = jobs_df[jobs_df["Country"].str.strip().str.lower() == cn.lower()] \
               if not jobs_df.empty else pd.DataFrame()
        vc   = 0
        if not rows.empty and v_col in rows.columns:
            vc = int((rows[v_col].str.lower() == "yes").sum())
        cells = dash.add_row().cells
        cells[0].text = cn
        cells[1].text = country["flag"]
        cells[2].text = str(len(rows))
        cells[3].text = f"✔ {vc}" if vc else "—"
        cells[4].text = country["resume"].split("/")[-1]
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if vc:
            _shade(cells[3], "E2EFDA")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ── All Jobs by Country ──────────────────────────────────────────
    _section(doc, f"2.  All {total} Job Listings by Country")
    gn = 1

    for country in TARGET_COUNTRIES:
        cn   = country["name"]
        rows = jobs_df[jobs_df["Country"].str.strip().str.lower() == cn.lower()] \
               if not jobs_df.empty else pd.DataFrame()
        if rows.empty:
            continue

        p = doc.add_paragraph()
        r = p.add_run(f"{country['flag']}  {cn}  —  {len(rows)} Jobs")
        r.bold = True; r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 70, 140)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)

        tbl = doc.add_table(rows=1, cols=7)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit   = False
        t_hdrs = ["#", "Job Title", "Company", "Location", "Skills", "Visa", "Apply"]
        t_wids = [Inches(0.3), Inches(2.1), Inches(1.6), Inches(1.2),
                  Inches(1.8), Inches(0.6), Inches(1.5)]
        for i, (h, w) in enumerate(zip(t_hdrs, t_wids)):
            _hdr_cell(tbl.rows[0].cells[i], h, bg="2F5597", width=w)

        for _, job in rows.iterrows():
            title  = str(job.get("Job Title", "N/A"))
            comp   = str(job.get("Company", "N/A"))
            loc    = str(job.get("Location", cn))
            skills = str(job.get("Required Skills", ""))[:60]
            url    = str(job.get("Job URL", ""))
            visa   = str(job.get(v_col, "No"))

            rc = tbl.add_row().cells
            rc[0].text = str(gn)
            rc[1].text = title
            rc[2].text = comp
            rc[3].text = loc
            rc[4].text = skills
            rc[5].text = visa
            rc[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if visa.lower() == "yes":
                _shade(rc[5], "C6EFCE")

            if url.startswith("http"):
                try:
                    _hyperlink(rc[6].paragraphs[0], url, "Apply →")
                except Exception:
                    rc[6].text = url[:40]
            else:
                rc[6].text = "N/A"

            rc[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            gn += 1

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Footer
    p = doc.add_paragraph()
    r = p.add_run(f"Generated by full_crawl_to_word.py  ·  {run_date}  ·  Hemachandiran Giri")
    r.font.size = Pt(8); r.font.color.rgb = RGBColor(150, 150, 150)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        doc.save(DOCX_PATH)
        log(f"Word document saved → {DOCX_PATH}", "OK")
    except PermissionError:
        alt = DOCX_PATH.parent / f"LinkedIn_FULL_{int(time.time())}.docx"
        doc.save(alt)
        log(f"File locked — saved to {alt}", "WARN")


# ══════════════════════════════════════════════════════════════════════
# PHASE 8 — Excel
# ══════════════════════════════════════════════════════════════════════

def save_excel(df: pd.DataFrame):
    if not HAS_EXCEL:
        return
    log("\nGenerating Excel …")
    wb = openpyxl.Workbook()
    ws = wb.active; ws.title = "All Jobs"

    cols = ["#","Date","Country","Flag","Tier","City","City Openings","Location","Company","Job Title",
            "Search Keyword","Posted Date","Easy Apply","Remote / Workplace","Match Score",
            "Visa Sponsorship","Skills","Resume","Applied","Notes","Job URL"]
    hf   = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    hfnt = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
    thin = Border(left=Side(style="thin"), right=Side(style="thin"),
                  top=Side(style="thin"),  bottom=Side(style="thin"))

    for ci, col in enumerate(cols, 1):
        c = ws.cell(row=1, column=ci, value=col)
        c.fill = hf; c.font = hfnt
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = thin

    vf   = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
    vcol = "Visa Sponsorship Mentioned"

    for ri, (_, job) in enumerate(df.iterrows(), 2):
        vv = str(job.get(vcol, "No"))
        vals = [ri-1, str(job.get("Crawl Date", job.get("Date", ""))), str(job.get("Country","")),
                str(job.get("Flag","")), str(job.get("Tier","")),
                str(job.get("City","")), str(job.get("City Openings","")),
                str(job.get("Location","")), str(job.get("Company","")),
                str(job.get("Job Title","")), str(job.get("Search Keyword","")),
                str(job.get("Posted Date","Past 3 weeks")),
                str(job.get("Easy Apply","No")),
                str(job.get("Remote / Workplace","On-site / Hybrid")),
                str(job.get("Match Score","85%")),
                vv, str(job.get("Required Skills","")),
                str(job.get("Resume File Path","")),
                str(job.get("Applied Status","No")),
                str(job.get("Notes","")),
                str(job.get("Job URL",""))]
        for ci, val in enumerate(vals, 1):
            c = ws.cell(row=ri, column=ci, value=val)
            c.font = Font(name="Calibri", size=9); c.border = thin
            if ci == 16 and vv.lower() == "yes":
                c.fill = vf

    for i, w in enumerate([4,12,14,5,5,16,14,20,22,32,20,16,12,18,12,15,30,45,8,20,55], 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"; ws.auto_filter.ref = ws.dimensions

    try:
        wb.save(MASTER_EXCEL_PATH)
        wb.save(EXCEL_PATH)
        log(f"Master Excel saved → {MASTER_EXCEL_PATH}", "OK")
        log(f"Run Excel archived → {EXCEL_PATH}", "OK")
    except PermissionError:
        alt = EXCEL_PATH.parent / f"full_crawl_{int(time.time())}.xlsx"
        wb.save(alt); log(f"Excel locked — saved to {alt}", "WARN")


# ══════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Full LinkedIn Crawler -> Word/Excel/CSV")
    parser.add_argument("--no-headless",     action="store_true",
                        help="Show browser window (required for first login or CAPTCHA)")
    parser.add_argument("--countries",       nargs="*", default=None,
                        help="Country names to scrape (e.g. netherlands ireland)")
    parser.add_argument("--max-per-keyword", type=int, default=999,
                        help="Max jobs per keyword per country (default: 999 = all)")
    parser.add_argument("--time-window", "--time-posted", type=str, default="24h",
                        help="Time window: 24h/1d (past 24 hrs), 1w/7d (past week), 1m/30d (past month), all")
    parser.add_argument("--skip-linkedin",   action="store_true")
    parser.add_argument("--skip-apis",       action="store_true")
    args = parser.parse_args()

    headless         = not args.no_headless
    countries_filter = [c.lower() for c in args.countries] if args.countries else None
    
    # Map time window argument to LinkedIn f_TPR parameter
    tw = (args.time_window or "24h").lower().strip()
    if tw in {"24h", "1d", "1day", "day"}:
        time_posted_range = "r86400"
        tw_label = "Past 24 Hours (1 Day)"
    elif tw in {"1w", "1week", "week", "7d"}:
        time_posted_range = "r604800"
        tw_label = "Past 1 Week (7 Days)"
    elif tw in {"1m", "1month", "month", "30d"}:
        time_posted_range = "r2592000"
        tw_label = "Past 1 Month (30 Days)"
    elif tw in {"all", "any"}:
        time_posted_range = ""
        tw_label = "All Time"
    else:
        time_posted_range = "r86400"
        tw_label = "Past 24 Hours (1 Day)"

    log("=" * 65)
    log("🚀  full_crawl_to_word.py  v2  (Voyager API + DOM fallback)")
    log(f"   Mode          : {'Headless' if headless else 'Visible browser'}")
    log(f"   Time Window   : {tw_label} ({time_posted_range or 'any'})")
    log(f"   Countries     : {countries_filter or 'All 10'}")
    log(f"   Max per kw    : {args.max_per_keyword}")
    log(f"   Output        : {WORK_DIR}")
    log(f"   Session       : {SESSION_DIR}")
    log("=" * 65)

    RUN_DIR.mkdir(parents=True, exist_ok=True)
    log(f"Run folder     : {RUN_DIR}")
    all_jobs = []

    if not args.skip_linkedin:
        all_jobs.extend(run_full_crawl(headless, countries_filter, args.max_per_keyword, time_posted_range))

    if not args.skip_apis:
        all_jobs.extend(fetch_free_apis())

    if not all_jobs:
        if MASTER_CSV_PATH.exists():
            log("No live data — loading existing Master CSV …", "WARN")
            jobs_df = pd.read_csv(MASTER_CSV_PATH, on_bad_lines="skip")
            jobs_df.dropna(subset=["Job URL"], inplace=True)
            sync_to_google_sheets(jobs_df)
        else:
            log("No data to process. Exiting.", "ERROR"); return
    else:
        jobs_df = save_csv(all_jobs)

    if jobs_df.empty:
        log("No data to write.", "ERROR"); return

    generate_word(jobs_df)
    save_excel(jobs_df)

    log("\n" + "=" * 65)
    log("✅  DONE", "OK")
    log(f"   Total unique jobs : {len(jobs_df)}")
    log(f"   📄 Word           : {DOCX_PATH}")
    log(f"   📊 Excel          : {EXCEL_PATH}")
    log(f"   📋 CSV            : {CSV_PATH}")
    log("=" * 65)

    for c in TARGET_COUNTRIES:
        cn  = c["name"]
        sub = jobs_df[jobs_df["Country"].str.strip().str.lower() == cn.lower()]
        vc  = 0
        vcol = "Visa Sponsorship Mentioned"
        if not sub.empty and vcol in sub.columns:
            vc = int((sub[vcol].str.lower() == "yes").sum())
        log(f"   {c['flag']}  {cn:<14}: {len(sub):>4} jobs  ({vc} mention visa/relocation)")


if __name__ == "__main__":
    main()
